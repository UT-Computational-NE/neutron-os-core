#!/usr/bin/env python3
"""
Markdown to Word Document Converter with Mermaid Diagram Support.

Converts markdown files to Word documents (.docx) with:
- Professional fonts (Calibri body, Aptos headings)
- Proper bullet list support
- Optimized Mermaid diagrams (full width, high quality)
- Minimal spacing for clean layout
- Cleaned of decorators and review notes
"""

import os
import re
import sys
import base64
import argparse
import json
import subprocess
from pathlib import Path
from io import BytesIO
from datetime import datetime
from typing import Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# DOCUMENT STYLING
# =============================================================================

def set_document_defaults(doc):
    """Set professional default fonts and spacing."""
    # Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # Reduce spacing after paragraphs
    pPr = normal_style._element.get_or_add_pPr()
    pSpacing = pPr.find(qn('w:spacing'))
    if pSpacing is None:
        pSpacing = OxmlElement('w:spacing')
        pPr.append(pSpacing)
    pSpacing.set(qn('w:after'), '60')
    pSpacing.set(qn('w:line'), '240')
    
    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Aptos'
        
        if i == 1:
            heading_style.font.size = Pt(18)
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)
        
        # Reduce heading spacing
        pPr = heading_style._element.get_or_add_pPr()
        pSpacing = pPr.find(qn('w:spacing'))
        if pSpacing is None:
            pSpacing = OxmlElement('w:spacing')
            pPr.append(pSpacing)
        
        pSpacing.set(qn('w:before'), '120')
        pSpacing.set(qn('w:after'), '60')


def set_paragraph_spacing(paragraph, space_after_pt: int = 6):
    """Set paragraph spacing in points."""
    pPr = paragraph._element.get_or_add_pPr()
    pSpacing = pPr.find(qn('w:spacing'))
    if pSpacing is None:
        pSpacing = OxmlElement('w:spacing')
        pPr.append(pSpacing)
    pSpacing.set(qn('w:after'), str(space_after_pt * 20))
    pSpacing.set(qn('w:line'), '240')


def set_list_restart(paragraph):
    """Set paragraph's list numbering to restart at 1."""
    pPr = paragraph._element.get_or_add_pPr()
    
    # Find existing numPr  
    numPr = pPr.find(qn('w:numPr'))
    
    if numPr is None:
        # Create new numPr if it doesn't exist
        numPr = OxmlElement('w:numPr')
        pPr.append(numPr)
        
        # Add numId element (required for numbering)
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(numId)
        
        # Add ilvl element (indentation level)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
    
    # Remove existing lvlRestart if present
    existing_restart = numPr.find(qn('w:lvlRestart'))
    if existing_restart is not None:
        numPr.remove(existing_restart)
    
    # Add lvlRestart element to restart numbering at 1
    # Note: This needs to be set on the first item of a new list
    restart = OxmlElement('w:lvlRestart')
    restart.set(qn('w:val'), '1')
    numPr.insert(0, restart)


def remove_reviewer_notes(content: str) -> str:
    """Remove reviewer notes, placeholders, and decorators."""
    # Remove blockquotes with review notes
    content = re.sub(
        r'>\s*\*\*Review Note:\*\*[^\n]*(?:\n>[^\n]*)*',
        '',
        content,
        flags=re.MULTILINE
    )
    
    # Remove placeholder patterns
    content = re.sub(r'\[PLACEHOLDER:[^\]]*\]', '', content)
    
    # Remove common emoji decorators
    emoji_pattern = r'[🚀📊🔐⚠️💡✅📝🔄🎯⚡🔥📈💎🎨🔧⚙️🧬✔️🎭🔮🏆🌟💫⭐]'
    content = re.sub(emoji_pattern, '', content)
    
    return content


def get_changed_files(src_path: Path) -> set[Path]:
    """Get markdown files changed in Git (modified, staged, or untracked)."""
    try:
        # Get the repository root
        repo_root = subprocess.check_output(
            ['git', 'rev-parse', '--show-toplevel'],
            cwd=str(src_path),
            stderr=subprocess.PIPE
        ).decode().strip()
        repo_root = Path(repo_root)
        
        # Get modified and staged files
        try:
            result = subprocess.check_output(
                ['git', 'diff', '--name-only', '--diff-filter=ACMRTUXB'],
                cwd=str(repo_root),
                stderr=subprocess.PIPE
            ).decode()
            modified = {repo_root / f for f in result.strip().split('\n') if f}
        except subprocess.CalledProcessError:
            modified = set()
        
        # Get staged files
        try:
            result = subprocess.check_output(
                ['git', 'diff', '--cached', '--name-only', '--diff-filter=ACMRTUXB'],
                cwd=str(repo_root),
                stderr=subprocess.PIPE
            ).decode()
            staged = {repo_root / f for f in result.strip().split('\n') if f}
        except subprocess.CalledProcessError:
            staged = set()
        
        # Get untracked files
        try:
            result = subprocess.check_output(
                ['git', 'ls-files', '--others', '--exclude-standard'],
                cwd=str(repo_root),
                stderr=subprocess.PIPE
            ).decode()
            untracked = {repo_root / f for f in result.strip().split('\n') if f}
        except subprocess.CalledProcessError:
            untracked = set()
        
        # Combine all changed files and filter to markdown files in src_path
        all_changed = modified | staged | untracked
        src_path_resolved = src_path.resolve()
        
        # Filter to markdown files that are within src_path
        md_changed = set()
        for f in all_changed:
            if f.suffix == '.md':
                try:
                    f.resolve().relative_to(src_path_resolved)
                    md_changed.add(f)
                except ValueError:
                    # File is not within src_path
                    pass
        
        if md_changed:
            print(f"Found {len(md_changed)} changed markdown file(s)\n", flush=True)
        
        return md_changed
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Not a git repo or git not available
        return set()


# =============================================================================
# MERMAID RENDERING
# =============================================================================

MAX_DIAGRAM_HEIGHT = 6.5  # inches (conservative limit to ensure diagrams fit with headers)
DIAGRAM_WIDTH = 6.5  # inches (page width 8.5" minus 1" margins on each side)
DIAGRAM_CACHE = {}  # In-memory cache for rendered diagrams
CACHE_FILE = Path(__file__).parent / '.diagram_cache.json'  # Persistent disk cache

def load_diagram_cache():
    """Load diagram cache from disk."""
    global DIAGRAM_CACHE
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, 'r') as f:
                cache_data = json.load(f)
                # Decode base64-encoded image data
                DIAGRAM_CACHE = {
                    code: base64.b64decode(img_b64)
                    for code, img_b64 in cache_data.items()
                }
            print(f"Loaded {len(DIAGRAM_CACHE)} cached diagrams", flush=True)
        except Exception as e:
            print(f"Warning: Could not load diagram cache: {e}", flush=True)
            DIAGRAM_CACHE = {}

def save_diagram_cache():
    """Save diagram cache to disk."""
    try:
        cache_data = {
            code: base64.b64encode(img_bytes).decode()
            for code, img_bytes in DIAGRAM_CACHE.items()
        }
        with open(CACHE_FILE, 'w') as f:
            json.dump(cache_data, f)
        print(f"Saved {len(DIAGRAM_CACHE)} diagrams to cache", flush=True)
    except Exception as e:
        print(f"Warning: Could not save diagram cache: {e}", flush=True)

def get_diagram_dimensions(image_data: bytes) -> Optional[tuple[float, float]]:
    """Get width and height of image in inches (assuming 96 DPI)."""
    from PIL import Image
    try:
        img = Image.open(BytesIO(image_data))
        # Mermaid renders at 96 DPI by default
        width_inches = img.width / 96
        height_inches = img.height / 96
        return (width_inches, height_inches)
    except Exception:
        return None


def render_mermaid_to_image(mermaid_code: str, timeout: int = 30, use_cache: bool = True) -> bytes | None:
    """Render Mermaid diagram to image bytes using mermaid.ink with optional caching.
    
    Sanitizes mermaid code to remove unsupported characters and fix common syntax issues.
    """
    import time
    
    # Fix common mermaid syntax issues
    sanitized_code = mermaid_code
    
    # Replace HTML line breaks with newlines inside quotes, spaces otherwise
    # This preserves multi-line labels in a way mermaid.ink can handle
    import re
    
    # Find all quoted strings and replace <br/> with space inside them
    # Mermaid.ink doesn't handle \n properly in labels, so use space instead
    def replace_br_in_quotes(match):
        content = match.group(0)
        content = content.replace('<br/>', ' ')
        content = content.replace('<br>', ' ')
        return content
    
    # Handle both single and double quoted strings
    sanitized_code = re.sub(r'"[^"]*"', replace_br_in_quotes, sanitized_code)
    sanitized_code = re.sub(r"'[^']*'", replace_br_in_quotes, sanitized_code)
    
    # For any remaining <br/> outside quotes, replace with space
    sanitized_code = sanitized_code.replace('<br/>', ' ')
    sanitized_code = sanitized_code.replace('<br>', ' ')
    
    # Replace emojis with text equivalents
    emoji_replacements = {
        '✅': 'ON',
        '⬜': 'OFF',
        '✔️': 'OK',
        '❌': 'X',
        '🥉': 'BRONZE',
        '🥈': 'SILVER', 
        '🥇': 'GOLD',
        '🔄': 'SYNC',
        '⚡': 'STREAM',
        '📊': 'CHART',
        '🚀': '',
        '📝': '',
        '🎯': '',
        '🔥': '',
        '📈': '',
        '💎': '',
        '🎨': '',
        '🔧': '',
        '⚙️': '',
        '🧬': '',
        '🎭': '',
        '🔮': '',
        '🏆': '',
        '🌟': '',
        '💫': '',
        '⭐': '',
    }
    
    for emoji, replacement in emoji_replacements.items():
        sanitized_code = sanitized_code.replace(emoji, replacement)
    
    # Handle xychart-beta (not supported by mermaid.ink)
    if 'xychart-beta' in sanitized_code:
        # Convert to a simple flowchart representation
        lines = sanitized_code.split('\n')
        title = ''
        for line in lines:
            if 'title' in line:
                title = line.split('"')[1] if '"' in line else 'Chart'
                break
        # Create a simple placeholder flowchart
        sanitized_code = f'graph LR\n    A[{title}]\n    B[Chart visualization not supported]\n    A --> B'
    
    # Remove any remaining Unicode emojis and special characters
    emoji_pattern = r'[\U0001F300-\U0001F9FF\U00002600-\U000027BF\U0001F900-\U0001FA9F\u2000-\u206F\u2190-\u21FF]'
    sanitized_code = re.sub(emoji_pattern, '', sanitized_code)
    
    # Clean up any problematic characters but keep necessary ASCII
    safe_chars = set(range(32, 127)) | {ord('\n'), ord('\r'), ord('\t')}
    sanitized_code = ''.join(char for char in sanitized_code if ord(char) in safe_chars)
    
    # Fix labels with square brackets - <br/> should already be replaced with spaces
    # Pattern for node definitions like NODE["label"]
    def fix_bracket_labels(text):
        pattern = r'(\w+)\["([^"]*)"\]'
        def replacer(match):
            node_id = match.group(1)
            label = match.group(2)
            # Keep the quotes in the output
            return f'{node_id}["{label}"]'
        return re.sub(pattern, replacer, text)
    
    sanitized_code = fix_bracket_labels(sanitized_code)
    
    # Also handle single-quoted labels - convert to double quotes  
    pattern2 = r"(\w+)\['([^']*)'\]"
    def single_quote_replacer(match):
        node_id = match.group(1)
        label = match.group(2)
        return f'{node_id}["{label}"]'
    sanitized_code = re.sub(pattern2, single_quote_replacer, sanitized_code)
    
    # Clean up any empty labels
    sanitized_code = re.sub(r'\[\s*\]', '[ ]', sanitized_code)
    
    # Check cache first (use original code as key for consistency)
    if use_cache and mermaid_code in DIAGRAM_CACHE:
        return DIAGRAM_CACHE[mermaid_code]
    
    try:
        # Encode sanitized diagram
        encoded = base64.urlsafe_b64encode(sanitized_code.encode()).decode()
        
        # Use width parameter for higher quality (scale only works with width/height)
        # 2400px width gives good quality when scaled down to 8 inches
        url = f"https://mermaid.ink/img/{encoded}?bgColor=white&width=2400"
        start = time.time()
        response = requests.get(url, timeout=timeout)
        elapsed = time.time() - start
        response.raise_for_status()
        
        image_data = response.content
        
        # Cache the result (with original code as key)
        if use_cache:
            DIAGRAM_CACHE[mermaid_code] = image_data
        
        return image_data
    except Exception as e:
        return None


def pregenerate_diagrams(blocks: list[dict], max_workers: int = 4) -> int:
    """Pregenerate all Mermaid diagrams asynchronously."""
    mermaid_blocks = [b for b in blocks if b['type'] == 'mermaid']
    
    if not mermaid_blocks:
        return 0
    
    print(f"\nPregenerating {len(mermaid_blocks)} diagram(s) asynchronously...", flush=True)
    
    rendered = 0
    failed = 0
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all rendering tasks
        futures = {
            executor.submit(render_mermaid_to_image, b['content']): i
            for i, b in enumerate(mermaid_blocks, 1)
        }
        
        # Collect results as they complete
        for future in as_completed(futures):
            idx = futures[future]
            try:
                image_data = future.result()
                if image_data:
                    rendered += 1
                    print(f"  ✓ Diagram {idx}/{len(mermaid_blocks)} rendered", flush=True)
                else:
                    failed += 1
                    print(f"  ✗ Diagram {idx}/{len(mermaid_blocks)} failed", flush=True)
            except Exception as e:
                failed += 1
                print(f"  ✗ Diagram {idx}/{len(mermaid_blocks)} error: {e}", flush=True)
    
    # Save cache to disk
    save_diagram_cache()
    print(f"Pregeneration complete: {rendered} rendered, {failed} failed\n", flush=True)
    
    return rendered


# =============================================================================
# MARKDOWN PARSING
# =============================================================================

def parse_markdown_to_blocks(md_content: str) -> list[dict]:
    """Parse markdown into structural blocks."""
    # Clean reviewer notes first
    md_content = remove_reviewer_notes(md_content)
    
    blocks = []
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Skip horizontal rules (---, ***, ___)
        if re.match(r'^\s*([-*_])\s*(\1\s*){2,}\s*$', line):
            i += 1
            continue
        
        # Heading
        if line.startswith('#'):
            match = re.match(r'^(#+)\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                content = match.group(2)
                blocks.append({'type': 'heading', 'level': level, 'content': content})
            i += 1
            continue
        
        # Mermaid diagram (check BEFORE generic code block)
        if line.strip().startswith('```mermaid'):
            diagram_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                diagram_lines.append(lines[i])
                i += 1
            if diagram_lines:
                blocks.append({'type': 'mermaid', 'content': '\n'.join(diagram_lines)})
            i += 1
            continue
        
        # Code block (generic, non-mermaid)
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                blocks.append({'type': 'code', 'content': '\n'.join(code_lines)})
            i += 1
            continue
        
        # Table (simple markdown table)
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'table', 'content': table_lines})
            continue
        
        # Blockquote
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                quote_lines.append(lines[i].lstrip('> ').strip())
                i += 1
            if quote_lines:
                blocks.append({'type': 'blockquote', 'content': '\n'.join(quote_lines)})
            continue
        
        # Unordered list
        if re.match(r'^[-*+]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[-*+]\s+', lines[i]):
                item = re.sub(r'^[-*+]\s+', '', lines[i])
                item = clean_inline_formatting(item.strip())
                if item:
                    items.append(item)
                i += 1
            if items:
                blocks.append({'type': 'list', 'list_type': 'bullet', 'content': items})
            continue
        
        # Ordered list
        if re.match(r'^\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i]):
                item = re.sub(r'^\d+\.\s+', '', lines[i])
                item = clean_inline_formatting(item.strip())
                if item:
                    items.append(item)
                i += 1
            if items:
                blocks.append({'type': 'list', 'list_type': 'number', 'content': items})
            continue
        
        # Regular paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip():
            # Stop if we hit any special markers
            if any(lines[i].startswith(p) for p in ['#', '```', '>', '|']):
                break
            # Stop if we hit a list marker at column 0
            if re.match(r'^([-*+]|\d+\.)\s+', lines[i]):
                break
            # Otherwise it's part of the paragraph
            para_lines.append(lines[i])
            i += 1
        
        if para_lines:
            para_text = ' '.join(para_lines).strip()
            para_text = clean_inline_formatting(para_text)
            if para_text:
                blocks.append({'type': 'paragraph', 'content': para_text})
    
    return blocks


def parse_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse markdown table."""
    if len(table_lines) < 2:
        return [], []
    
    def parse_row(line: str) -> list[str]:
        cells = line.split('|')[1:-1]
        return [clean_inline_formatting(c.strip()) for c in cells]
    
    headers = parse_row(table_lines[0])
    rows = [
        parse_row(line)
        for line in table_lines[2:]
        if line.strip() and not re.match(r'^[\s|:-]+$', line)
    ]
    
    return headers, rows


def apply_inline_formatting(paragraph, text: str):
    """Apply markdown inline formatting to a Word paragraph (bold, italic, code, links)."""
    # Pattern: **bold**, __bold__, *italic*, _italic_, `code`, [link](url)
    pattern = r'(\*\*(.+?)\*\*|__(.+?)__|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|_(.+?)_|`(.+?)`|\[(.+?)\]\((.+?)\))'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        matched_text = match.group(0)
        
        # Bold: **text** or __text__
        if matched_text.startswith('**') or matched_text.startswith('__'):
            content = match.group(2) or match.group(3)
            run = paragraph.add_run(content)
            run.bold = True
        
        # Italic: *text* or _text_ (but not ** or __)
        elif matched_text.startswith('*') or matched_text.startswith('_'):
            content = match.group(4) or match.group(5)
            run = paragraph.add_run(content)
            run.italic = True
        
        # Code: `text`
        elif matched_text.startswith('`'):
            content = match.group(6)
            run = paragraph.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        # Link: [text](url)
        elif matched_text.startswith('['):
            content = match.group(7)
            url = match.group(8)
            run = paragraph.add_run(content)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        
        last_end = match.end()
    
    # Add remaining text after last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def clean_inline_formatting(text: str) -> str:
    """Remove markdown formatting while preserving content (for simple text extraction)."""
    # Bold and italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Code backticks
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Links
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text.strip()


# =============================================================================
# DOCUMENT GENERATION
# =============================================================================

def add_table_to_doc(doc, headers: list[str], rows: list[list[str]]):
    """Add a formatted table."""
    if not headers:
        return
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            set_paragraph_spacing(paragraph, 6)
    
    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row):
                row[i].text = cell_data
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                    set_paragraph_spacing(paragraph, 4)
    
    # Space after table
    doc.add_paragraph()


def add_mermaid_to_doc(doc, mermaid_code: str, keep_with_previous: bool = True):
    """Add a Mermaid diagram optimized for Word with size validation.
    
    Diagrams are sized to fit within document margins and centered properly.
    The keep_with_previous parameter helps prevent separation from headers.
    """
    image_data = render_mermaid_to_image(mermaid_code)
    
    if image_data:
        image_stream = BytesIO(image_data)
        try:
            # Get dimensions for validation and scaling
            dims = get_diagram_dimensions(image_data)
            
            # Use the document's actual margins to calculate available width
            # Standard Word margins are 1" on each side
            # We'll use 6.5" as the safe maximum width within margins
            max_width = Inches(6.5)  # 8.5" page - 1" left - 1" right margins
            max_height = Inches(MAX_DIAGRAM_HEIGHT)  # Maximum height to fit on page
            
            # Calculate actual dimensions to use
            final_width = max_width
            final_height = None  # Let Word auto-scale by default
            
            if dims:
                width_inches, height_inches = dims
                
                # Calculate aspect ratio
                aspect_ratio = width_inches / height_inches
                
                # Check if we need to scale down due to height constraint
                if height_inches > MAX_DIAGRAM_HEIGHT:
                    print(f"    [Scaling: Diagram height {height_inches:.1f}\" exceeds {MAX_DIAGRAM_HEIGHT}\" limit - scaling down]", flush=True)
                    # Scale down to fit height constraint
                    final_height = max_height
                    final_width = Inches(MAX_DIAGRAM_HEIGHT * aspect_ratio)
                    
                    # But also ensure width doesn't exceed page width
                    if final_width > max_width:
                        final_width = max_width
                        final_height = Inches(6.5 / aspect_ratio)
                
                # Also check if the natural width at 6.5" would make it too tall
                else:
                    potential_height = 6.5 / aspect_ratio
                    if potential_height > MAX_DIAGRAM_HEIGHT:
                        # Scale to fit height instead
                        final_height = max_height
                        final_width = Inches(MAX_DIAGRAM_HEIGHT * aspect_ratio)
            
            # Add paragraph for diagram
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Always keep with previous paragraph to prevent separation from headers
            # This is especially important when a diagram follows immediately after a section heading
            pPr = p._element.get_or_add_pPr()
            # Check if keepPrev already exists before adding
            if not pPr.find(qn('w:keepPrev')):
                keepWithPrev = OxmlElement('w:keepPrev')
                pPr.append(keepWithPrev)
            
            # Also add keep-next to the previous paragraph if it's a heading
            # This provides double insurance against separation
            if doc.paragraphs and len(doc.paragraphs) > 1:
                prev_para = doc.paragraphs[-2]  # Get the paragraph before the one we just added
                if prev_para.style and 'Heading' in prev_para.style.name:
                    prev_pPr = prev_para._element.get_or_add_pPr()
                    # Check if keepNext already exists before adding
                    if not prev_pPr.find(qn('w:keepNext')):
                        keepNext = OxmlElement('w:keepNext')
                        prev_pPr.append(keepNext)
            
            # Add image with calculated dimensions
            run = p.add_run()
            if final_height:
                run.add_picture(image_stream, width=final_width, height=final_height)
            else:
                run.add_picture(image_stream, width=final_width)
            
            set_paragraph_spacing(p, 6)
        except Exception as e:
            print(f"    [Error: Failed to add diagram: {e}]", flush=True)
            p = doc.add_paragraph("[Diagram failed to render]")
            set_paragraph_spacing(p, 6)
    else:
        # Provide more helpful fallback with diagram type info
        diagram_type = "diagram"
        if "flowchart" in mermaid_code.lower()[:20]:
            diagram_type = "flowchart"
        elif "xychart" in mermaid_code.lower()[:20]:
            diagram_type = "chart"
        elif "gantt" in mermaid_code.lower()[:20]:
            diagram_type = "Gantt chart"
        elif "sequence" in mermaid_code.lower()[:20]:
            diagram_type = "sequence diagram"
        
        p = doc.add_paragraph(f"[{diagram_type.title()} - See original markdown for details]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a note in italics
        run = p.runs[0]
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        set_paragraph_spacing(p, 6)


def add_paragraph_to_doc(doc, text: str, style: Optional[str] = None):
    """Add a paragraph with proper spacing and inline formatting."""
    p = doc.add_paragraph(style=style)
    apply_inline_formatting(p, text)
    if p.style:
        p.style.font.name = 'Calibri'
    set_paragraph_spacing(p, 6)
    return p


def convert_md_to_docx(
    md_path,
    output_path: Optional[str] = None,
    add_draft_notice: bool = False,
    pregenerate: bool = False
) -> str:
    """Convert markdown to Word document."""
    md_path = Path(md_path)
    
    if not md_path.exists():
        raise FileNotFoundError(f"File not found: {md_path}")
    
    if output_path is None:
        output_path = md_path.with_suffix('.docx')  # type: ignore
    else:
        output_path = Path(output_path)  # type: ignore
    
    # Ensure output_path is a Path object
    assert isinstance(output_path, Path)
    
    # Read and parse
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    blocks = parse_markdown_to_blocks(content)
    
    # Pregenerate diagrams if requested or if any diagrams exist
    if pregenerate or any(b['type'] == 'mermaid' for b in blocks):
        pregenerate_diagrams(blocks)
    
    # Create document
    doc = Document()
    set_document_defaults(doc)
    
    # Add draft notice if requested
    if add_draft_notice:
        p = doc.add_paragraph()
        run = p.add_run('DRAFT - FOR INTERNAL REVIEW ONLY')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(192, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, 0)
        
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(128, 128, 128)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, 12)
    
    # Add content blocks
    mermaid_count = 0
    previous_type = None
    for i, block in enumerate(blocks, 1):
        block_type = block['type']
        
        if block_type == 'heading':
            level = min(block['level'], 9)
            heading = doc.add_heading(block['content'], level=level)
            # Always keep headings with next element to prevent orphaned headings
            # This is especially important for diagrams, tables, and lists
            if i < len(blocks):
                next_type = blocks[i]['type']
                # Always apply keep-next for headings followed by content
                if next_type in ['mermaid', 'table', 'list', 'code', 'paragraph']:
                    heading_pPr = heading._element.get_or_add_pPr()
                    # Check if keepNext already exists before adding
                    if not heading_pPr.find(qn('w:keepNext')):
                        keepNext = OxmlElement('w:keepNext')
                        heading_pPr.append(keepNext)
            print(f"  [{i}/{len(blocks)}] Heading", flush=True)
        
        elif block_type == 'paragraph':
            add_paragraph_to_doc(doc, block['content'])
            print(f"  [{i}/{len(blocks)}] Paragraph", flush=True)
        
        elif block_type == 'mermaid':
            mermaid_count += 1
            print(f"  [{i}/{len(blocks)}] Diagram {mermaid_count}...", end=" ", flush=True)
            # Always keep diagrams with previous content to prevent orphaning
            # This is especially important after headings, but also helps with flow
            keep_with_prev = True
            add_mermaid_to_doc(doc, block['content'], keep_with_previous=keep_with_prev)
        
        elif block_type == 'code':
            p = doc.add_paragraph()
            run = p.add_run(block['content'])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            set_paragraph_spacing(p, 6)
            print(f"  [{i}/{len(blocks)}] Code block", flush=True)
        
        elif block_type == 'table':
            headers, rows = parse_table(block['content'])
            add_table_to_doc(doc, headers, rows)
            print(f"  [{i}/{len(blocks)}] Table", flush=True)
        
        elif block_type == 'list':
            for idx, item in enumerate(block['content']):
                if block['list_type'] == 'number':
                    p = doc.add_paragraph(item, style='List Number')
                    if idx == 0:  # First item in this ordered list - restart numbering
                        set_list_restart(p)
                else:
                    p = doc.add_paragraph(item, style='List Bullet')
                set_paragraph_spacing(p, 3)
            print(f"  [{i}/{len(blocks)}] List", flush=True)
        
        elif block_type == 'blockquote':
            p = doc.add_paragraph()
            for run in p.runs:  # Clear any default runs
                p._element.remove(run._element)
            apply_inline_formatting(p, block['content'])
            run = p.runs[0] if p.runs else p.add_run()
            run.italic = True
            p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else None
            set_paragraph_spacing(p, 6)
            print(f"  [{i}/{len(blocks)}] Quote", flush=True)
        
        # Track previous block type for context
        previous_type = block_type
    
    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    print(f"  Saving document...", end=" ", flush=True)
    doc.save(str(output_path))
    print(f"Done!", flush=True)
    
    print(f"\nGenerated: {output_path.name}")
    if mermaid_count > 0:
        print(f"  Processed {len(blocks)} blocks ({mermaid_count} diagram{'s' if mermaid_count > 1 else ''})")
    
    return str(output_path)


def convert_all_docs(
    src_dir,
    output_dir: Optional[str] = None,
    skip_existing: bool = True,
    add_draft_notice: bool = False,
    pregenerate: bool = False,
    changed_only: bool = False
) -> list[str]:
    """Convert all .md files in directory (recursively) to .docx.
    
    Args:
        src_dir: Source directory with markdown files
        output_dir: Output directory (default: src_dir/generated)
        skip_existing: Skip if .docx already exists
        add_draft_notice: Add draft watermark
        pregenerate: Pregenerate all diagrams asynchronously
        changed_only: Only convert files changed in Git
    """
    src_path = Path(src_dir)
    
    if not src_path.is_dir():
        raise NotADirectoryError(f"Directory not found: {src_dir}")
    
    # Determine output directory
    if output_dir is None:
        output_dir = str(src_path / 'generated')  # type: ignore
    
    output_dir_path = Path(output_dir)  # type: ignore
    output_dir_path.mkdir(parents=True, exist_ok=True)
    
    # Find markdown files
    md_files = sorted(src_path.glob('**/*.md'))
    
    if not md_files:
        print(f"No markdown files found in {src_path}")
        return []
    
    # Filter to changed files if requested
    if changed_only:
        changed_files = get_changed_files(src_path)
        if not changed_files:
            print("No changed markdown files found in Git")
            return []
        # Resolve md_files to absolute paths for comparison
        md_files_resolved = {f.resolve() for f in md_files}
        changed_resolved = {f.resolve() for f in changed_files}
        md_files = sorted([f for f in md_files if f.resolve() in changed_resolved])
        if not md_files:
            print(f"No changed markdown files found in {src_path}")
            return []
    
    # If pregenerating, parse all files first and collect all diagrams
    if pregenerate:
        print(f"Collecting diagrams from {len(md_files)} file(s)...")
        all_blocks = []
        for md_file in md_files:
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            all_blocks.extend(parse_markdown_to_blocks(content))
        pregenerate_diagrams(all_blocks)
    
    print(f"\nConverting {len(md_files)} markdown file(s)...")
    print(f"Output: {output_dir_path}\n")
    
    generated = []
    for i, md_file in enumerate(md_files, 1):
        rel_path = md_file.relative_to(src_path)
        print(f"[{i}/{len(md_files)}] {rel_path}")
        
        # Maintain same filename, just change extension
        output_path = output_dir_path / md_file.stem
        output_path = output_path.with_suffix('.docx')
        
        if skip_existing and output_path.exists():
            print(f"    Skipped (exists)\n")
            generated.append(str(output_path))
            continue
        
        try:
            result = convert_md_to_docx(
                str(md_file),
                str(output_path),
                add_draft_notice=add_draft_notice,
                pregenerate=False  # Already done above if requested
            )
            generated.append(result)
            print()
        except Exception as e:
            print(f"    ERROR: {e}\n")
    
    print(f"Completed: {len(generated)} file(s)")
    return generated


def main():
    """CLI entry point."""
    # Load cached diagrams at startup
    load_diagram_cache()
    
    parser = argparse.ArgumentParser(
        description="Convert markdown to professional Word documents with Mermaid support"
    )
    parser.add_argument('source', help='Markdown file or directory')
    parser.add_argument('output', nargs='?', help='Output file or directory')
    parser.add_argument('--all', action='store_true', help='Convert all .md files recursively')
    parser.add_argument('--draft', action='store_true', help='Add draft notice')
    parser.add_argument('--skip', action='store_true', help='Skip existing files')
    parser.add_argument('--changed', action='store_true', help='Only convert files changed in Git (requires --all)')
    parser.add_argument('--pregenerate', action='store_true', help='Pregenerate all diagrams asynchronously before conversion')
    parser.add_argument('--check-diagrams', action='store_true', help='Only validate diagram sizes without generating Word docs')
    
    args = parser.parse_args()
    
    if args.check_diagrams:
        # Check diagram sizes only
        source_path = Path(args.source)
        if source_path.is_file():
            print(f"Checking diagrams in {source_path.name}...\n")
            with open(source_path, 'r', encoding='utf-8') as f:
                content = f.read()
            blocks = parse_markdown_to_blocks(content)
            
            diagram_count = 0
            oversized_count = 0
            for block in blocks:
                if block['type'] == 'mermaid':
                    diagram_count += 1
                    image_data = render_mermaid_to_image(block['content'])
                    if image_data:
                        dims = get_diagram_dimensions(image_data)
                        if dims:
                            width, height = dims
                            status = "OK" if height <= MAX_DIAGRAM_HEIGHT else "OVERSIZED"
                            print(f"  Diagram {diagram_count}: {height:.1f}\" height ({width:.1f}\" wide) [{status}]")
                            if height > MAX_DIAGRAM_HEIGHT:
                                oversized_count += 1
                                print(f"    -> Exceeds {MAX_DIAGRAM_HEIGHT}\" limit by {height - MAX_DIAGRAM_HEIGHT:.1f}\"")
            
            print()
            if oversized_count == 0:
                print(f"✓ All {diagram_count} diagram(s) fit within size limits")
            else:
                print(f"✗ {oversized_count}/{diagram_count} diagram(s) exceed size limits")
                print("\n  Fix these diagrams at the source before converting to Word.")
    elif args.all:
        convert_all_docs(
            args.source,
            args.output,
            skip_existing=args.skip,
            add_draft_notice=args.draft,
            pregenerate=args.pregenerate,
            changed_only=args.changed
        )
    else:
        convert_md_to_docx(args.source, args.output, add_draft_notice=args.draft, pregenerate=args.pregenerate)


if __name__ == '__main__':
    main()
