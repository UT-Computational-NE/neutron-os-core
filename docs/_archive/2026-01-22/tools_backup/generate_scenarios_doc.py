#!/usr/bin/env python3
"""
Generate Word document with proposed Superset scenarios for Nick's review.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = str(cell_data)
    
    return table

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Neutron OS Data Platform', 0)
    subtitle = doc.add_paragraph('Proposed Superset Dashboard Scenarios')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Status: DRAFT - For Review')
    doc.add_paragraph('Reviewer: Nick Luciano')
    doc.add_paragraph()
    
    # Introduction
    add_heading(doc, 'Introduction', 1)
    doc.add_paragraph(
        'This document outlines proposed Superset dashboard scenarios for the Neutron OS data platform. '
        'These scenarios will drive the design of our data lakehouse architecture using a test-driven approach: '
        'we define what users need to see first, then build the data pipelines to support those visualizations.'
    )
    doc.add_paragraph(
        'Please review each scenario and provide feedback on:'
    )
    bullets = doc.add_paragraph()
    bullets.add_run('• ').bold = True
    bullets.add_run('Questions/metrics that are missing or should be prioritized\n')
    bullets.add_run('• ').bold = True
    bullets.add_run('Charts or visualizations that would be most valuable\n')
    bullets.add_run('• ').bold = True
    bullets.add_run('Filters needed for practical use\n')
    bullets.add_run('• ').bold = True
    bullets.add_run('Data sources or time ranges to consider\n')
    bullets.add_run('• ').bold = True
    bullets.add_run('Priority order for implementation')
    
    doc.add_page_break()
    
    # ===========================================
    # Scenario 1: Reactor Operations Dashboard
    # ===========================================
    add_heading(doc, 'Scenario 1: Reactor Operations Dashboard', 1)
    add_heading(doc, 'Priority: HIGH (First to implement)', 2)
    
    add_heading(doc, 'User Story', 3)
    doc.add_paragraph(
        'As a reactor operator or facility manager, I want to see the current and recent state of the reactor '
        'so that I can monitor operations, identify trends, and respond to anomalies.'
    )
    
    add_heading(doc, 'Questions This Dashboard Answers', 3)
    questions = [
        'What is the current reactor power and how has it changed today?',
        'What are the current control rod positions?',
        'What are the fuel and pool water temperatures?',
        'Are there any anomalies or unexpected readings?',
        'What was the power history over the last shift/day/week?',
    ]
    for q in questions:
        doc.add_paragraph(f'• {q}', style='List Bullet')
    
    add_heading(doc, 'Proposed Charts', 3)
    add_table(doc, 
        ['Chart Name', 'Type', 'Description', 'Priority'],
        [
            ['Power Timeline', 'Line chart', 'Real-time power (kW) over selected time range', 'P0'],
            ['Rod Positions', 'Multi-line or bar', 'Current position of Tran, Shim1, Shim2, Reg rods', 'P0'],
            ['Temperature Gauges', 'Gauge/KPI cards', 'Current fuel temp and water temp with thresholds', 'P0'],
            ['Power Distribution', 'Histogram', 'Distribution of power levels over time period', 'P1'],
            ['Detector Signals', 'Line chart', 'NM, NPP, NP signals (for advanced users)', 'P2'],
        ]
    )
    
    add_heading(doc, 'Proposed Filters', 3)
    add_table(doc,
        ['Filter', 'Type', 'Default', 'Notes'],
        [
            ['Date Range', 'Date picker', 'Last 24 hours', 'Quick presets: Today, Yesterday, Last 7 days'],
            ['Power Threshold', 'Slider', '> 0 kW', 'Filter out zero-power periods'],
            ['Auto-refresh', 'Toggle', 'On (5 min)', 'For real-time monitoring'],
        ]
    )
    
    add_heading(doc, 'Data Sources', 3)
    doc.add_paragraph('• serial_data/*.csv - Reactor time-series (power, temps, rod positions)')
    doc.add_paragraph('• Refresh: Near real-time (aim for 5-minute latency)')
    
    add_heading(doc, 'Questions for Nick', 3)
    doc.add_paragraph('[   ] What time resolution is needed? (per-second, per-minute, hourly?)', style='List Bullet')
    doc.add_paragraph('[   ] Which metrics are most critical for at-a-glance monitoring?', style='List Bullet')
    doc.add_paragraph('[   ] Should there be alert thresholds highlighted on charts?', style='List Bullet')
    doc.add_paragraph('[   ] What historical range is typically needed? (hours, days, weeks?)', style='List Bullet')
    
    doc.add_page_break()
    
    # ===========================================
    # Scenario 2: Reactor Performance Analytics
    # ===========================================
    add_heading(doc, 'Scenario 2: Reactor Performance Analytics', 1)
    add_heading(doc, 'Priority: HIGH', 2)
    
    add_heading(doc, 'User Story', 3)
    doc.add_paragraph(
        'As a researcher or operator, I want to analyze reactor performance over time by correlating power output, '
        'xenon poisoning, fuel burnup, and control rod positions so that I can understand operational patterns, '
        'optimize startup procedures, and predict reactivity requirements.'
    )
    
    add_heading(doc, 'Questions This Dashboard Answers', 3)
    questions = [
        'How does xenon concentration correlate with power history?',
        'What is the current excess reactivity given fuel burnup and xenon state?',
        'How have control rod positions changed relative to power demand?',
        'Which fuel elements have the highest burnup?',
        'What is the typical startup time from cold critical to full power?',
    ]
    for q in questions:
        doc.add_paragraph(f'• {q}', style='List Bullet')
    
    add_heading(doc, 'Proposed Charts', 3)
    add_table(doc, 
        ['Chart Name', 'Type', 'Description', 'Priority'],
        [
            ['Power & Xenon Timeline', 'Dual-axis line', 'Power (kW) and Xe-135 concentration over time', 'P0'],
            ['Rod Position vs Power', 'Scatter plot', 'Correlation between power and avg rod height, colored by rod', 'P0'],
            ['Fuel Burnup Heatmap', 'Heatmap', 'Core layout showing U-235 burnup by position (hexagonal)', 'P1'],
            ['Excess Reactivity Trend', 'Line chart', 'Calculated excess reactivity over time', 'P1'],
            ['Temperature Correlation', 'Scatter plot', 'Fuel temp vs water temp, colored by power', 'P2'],
            ['Startup Time Distribution', 'Histogram', 'Time to reach full power from cold critical', 'P2'],
            ['Daily Energy Production', 'Bar chart', 'Integrated energy (MWh) by day', 'P1'],
        ]
    )
    
    add_heading(doc, 'Proposed Filters', 3)
    add_table(doc,
        ['Filter', 'Type', 'Default', 'Notes'],
        [
            ['Date Range', 'Date picker', 'Last 7 days', ''],
            ['Core Configuration', 'Dropdown', 'Latest', 'Select BOC/EOC configuration'],
            ['Power Threshold', 'Slider', '> 0 kW', 'Filter operational periods only'],
            ['Rod Selection', 'Multi-select', 'All', 'Filter specific rods'],
        ]
    )
    
    add_heading(doc, 'Data Sources', 3)
    doc.add_paragraph('• serial_data/*.csv - Reactor time-series')
    doc.add_paragraph('• Xe_burnup_2025.csv - Xenon/Iodine dynamics (433K+ rows)')
    doc.add_paragraph('• static/core/*.csv - Core configurations with burnup')
    doc.add_paragraph('• rho_vs_T.csv - Reactivity vs temperature lookup')
    
    add_heading(doc, 'Questions for Nick', 3)
    doc.add_paragraph('[   ] Is the Xenon correlation the most valuable insight here?', style='List Bullet')
    doc.add_paragraph('[   ] Should MPACT shadow predictions be overlaid on measured data?', style='List Bullet')
    doc.add_paragraph('[   ] What burnup thresholds would trigger attention? (for heatmap colors)', style='List Bullet')
    doc.add_paragraph('[   ] How is "startup time" currently defined/measured?', style='List Bullet')
    
    doc.add_page_break()
    
    # ===========================================
    # Scenario 3: Elog Activity Summary
    # ===========================================
    add_heading(doc, 'Scenario 3: Elog Activity Summary', 1)
    add_heading(doc, 'Priority: MEDIUM', 2)
    
    add_heading(doc, 'User Story', 3)
    doc.add_paragraph(
        'As a facility manager or regulatory inspector, I want to see a summary of operations log activity '
        'so that I can understand operational patterns, verify compliance, and prepare for audits.'
    )
    
    add_heading(doc, 'Questions This Dashboard Answers', 3)
    questions = [
        'How many log entries were created per day/week/month?',
        'Which operators have logged the most entries?',
        'What types of operations are most frequently logged?',
        'Are there gaps in logging that need investigation?',
        'What is the run history over a given period?',
    ]
    for q in questions:
        doc.add_paragraph(f'• {q}', style='List Bullet')
    
    add_heading(doc, 'Proposed Charts', 3)
    add_table(doc, 
        ['Chart Name', 'Type', 'Description', 'Priority'],
        [
            ['Entries Per Day', 'Bar chart', 'Count of log entries by day', 'P0'],
            ['Entries by Operator', 'Pie/donut', 'Distribution of entries across operators', 'P1'],
            ['Run Timeline', 'Gantt/timeline', 'Visual timeline of reactor runs', 'P1'],
            ['Entry Categories', 'Bar chart', 'Breakdown by entry type (startup, shutdown, observation)', 'P2'],
            ['Logging Gaps', 'Calendar heatmap', 'Highlight days with few/no entries', 'P2'],
        ]
    )
    
    add_heading(doc, 'Proposed Filters', 3)
    add_table(doc,
        ['Filter', 'Type', 'Default', 'Notes'],
        [
            ['Date Range', 'Date picker', 'Last 30 days', ''],
            ['Operator', 'Multi-select', 'All', ''],
            ['Run Number', 'Dropdown', 'All', ''],
            ['Keywords', 'Text search', '', 'Full-text search across entries'],
        ]
    )
    
    add_heading(doc, 'Data Sources', 3)
    doc.add_paragraph('• Elog system (currently JSON files, migrating to immutable blockchain)')
    doc.add_paragraph('• Note: This dashboard depends on elog system development (see Elog PRD)')
    
    add_heading(doc, 'Questions for Nick', 3)
    doc.add_paragraph('[   ] What categories/tags should elog entries have?', style='List Bullet')
    doc.add_paragraph('[   ] What constitutes a "gap" that should be flagged?', style='List Bullet')
    doc.add_paragraph('[   ] Should this include export-to-PDF for audit preparation?', style='List Bullet')
    
    doc.add_page_break()
    
    # ===========================================
    # Scenario 4: Experiment Tracking
    # ===========================================
    add_heading(doc, 'Scenario 4: Experiment Tracking', 1)
    add_heading(doc, 'Priority: MEDIUM', 2)
    
    add_heading(doc, 'User Story', 3)
    doc.add_paragraph(
        'As a researcher or principal investigator, I want to track experiments from planning through completion '
        'so that I can manage research activities, correlate results with reactor conditions, and report on progress.'
    )
    
    add_heading(doc, 'Questions This Dashboard Answers', 3)
    questions = [
        'What experiments are currently planned, in progress, or completed?',
        'What reactor conditions were present during each experiment?',
        'How much beam time has each PI/project used?',
        'What is the backlog of experiments awaiting scheduling?',
        'Which experiments correlate with specific reactor runs?',
    ]
    for q in questions:
        doc.add_paragraph(f'• {q}', style='List Bullet')
    
    add_heading(doc, 'Proposed Charts', 3)
    add_table(doc, 
        ['Chart Name', 'Type', 'Description', 'Priority'],
        [
            ['Experiment Status Funnel', 'Funnel', 'Count by status: Planned → Scheduled → Running → Completed', 'P0'],
            ['Experiment Calendar', 'Calendar', 'Scheduled experiments on calendar view', 'P1'],
            ['Irradiation Hours by PI', 'Bar chart', 'Usage breakdown by principal investigator', 'P1'],
            ['Experiment Timeline', 'Gantt', 'Timeline showing experiment duration and status', 'P2'],
        ]
    )
    
    add_heading(doc, 'Proposed Filters', 3)
    add_table(doc,
        ['Filter', 'Type', 'Default', 'Notes'],
        [
            ['Date Range', 'Date picker', 'Current quarter', ''],
            ['Status', 'Multi-select', 'All', 'Planned, Scheduled, Running, Completed, Cancelled'],
            ['Principal Investigator', 'Dropdown', 'All', ''],
            ['Experiment Type', 'Multi-select', 'All', ''],
        ]
    )
    
    add_heading(doc, 'Data Sources', 3)
    doc.add_paragraph('• Experiment tracking system (to be developed)')
    doc.add_paragraph('• Correlation with reactor time-series for conditions during experiments')
    doc.add_paragraph('• Note: This depends on experiment management features (future development)')
    
    add_heading(doc, 'Questions for Nick', 3)
    doc.add_paragraph('[   ] How are experiments currently tracked? (spreadsheet, system, ad-hoc?)', style='List Bullet')
    doc.add_paragraph('[   ] What experiment metadata is most important?', style='List Bullet')
    doc.add_paragraph('[   ] Should this integrate with a scheduling/calendar system?', style='List Bullet')
    
    doc.add_page_break()
    
    # ===========================================
    # Scenario 5: Audit Readiness
    # ===========================================
    add_heading(doc, 'Scenario 5: Audit Readiness', 1)
    add_heading(doc, 'Priority: MEDIUM-HIGH', 2)
    
    add_heading(doc, 'User Story', 3)
    doc.add_paragraph(
        'As a regulatory inspector or compliance officer, I want to verify the integrity of historical records '
        'and generate evidence packages so that I can conduct audits efficiently and with confidence.'
    )
    
    add_heading(doc, 'Questions This Dashboard Answers', 3)
    questions = [
        'Can I verify that records have not been tampered with?',
        'What audit events have occurred in a given period?',
        'What evidence packages are available for inspection?',
        'Are there any verification failures or anomalies?',
        'What is the complete audit trail for a specific record?',
    ]
    for q in questions:
        doc.add_paragraph(f'• {q}', style='List Bullet')
    
    add_heading(doc, 'Proposed Charts', 3)
    add_table(doc, 
        ['Chart Name', 'Type', 'Description', 'Priority'],
        [
            ['Audit Event Timeline', 'Timeline', 'All audit events (data changes, access, exports)', 'P0'],
            ['Verification Status', 'KPI cards', 'Count of verified, pending, failed records', 'P0'],
            ['Evidence Package Inventory', 'Table', 'List of available audit packages with download links', 'P1'],
            ['Data Integrity Check', 'Status indicator', 'Visual confirmation of blockchain/Merkle verification', 'P1'],
        ]
    )
    
    add_heading(doc, 'Proposed Filters', 3)
    add_table(doc,
        ['Filter', 'Type', 'Default', 'Notes'],
        [
            ['Date Range', 'Date picker', 'Last 90 days', ''],
            ['Event Type', 'Multi-select', 'All', 'Create, Update, Access, Export, Verify'],
            ['Record Type', 'Multi-select', 'All', 'Elog, Reactor Data, Experiment, etc.'],
            ['Verification Status', 'Dropdown', 'All', 'Verified, Pending, Failed'],
        ]
    )
    
    add_heading(doc, 'Data Sources', 3)
    doc.add_paragraph('• Hyperledger Fabric blockchain (audit events, Merkle proofs)')
    doc.add_paragraph('• Immudb (single-facility audit trail during development)')
    doc.add_paragraph('• Note: This depends on blockchain infrastructure development')
    
    add_heading(doc, 'Questions for Nick', 3)
    doc.add_paragraph('[   ] What does a typical NRC inspection request look like?', style='List Bullet')
    doc.add_paragraph('[   ] What format should evidence packages be in? (PDF, ZIP, other?)', style='List Bullet')
    doc.add_paragraph('[   ] How far back do inspectors typically need to query?', style='List Bullet')
    
    doc.add_page_break()
    
    # ===========================================
    # Summary & Next Steps
    # ===========================================
    add_heading(doc, 'Summary: Proposed Priority Order', 1)
    
    add_table(doc,
        ['Priority', 'Scenario', 'Rationale'],
        [
            ['1', 'Reactor Operations Dashboard', 'Most immediate value; validates core time-series pipeline'],
            ['2', 'Reactor Performance Analytics', 'Combines multiple data sources; validates join logic'],
            ['3', 'Audit Readiness', 'Critical for compliance; drives blockchain requirements'],
            ['4', 'Elog Activity Summary', 'Depends on elog system development'],
            ['5', 'Experiment Tracking', 'Depends on experiment management system'],
        ]
    )
    
    doc.add_paragraph()
    
    add_heading(doc, 'Next Steps', 2)
    doc.add_paragraph('1. Nick reviews and provides feedback on this document')
    doc.add_paragraph('2. Prioritize and refine scenarios based on feedback')
    doc.add_paragraph('3. Define Gold table schemas for highest-priority scenario')
    doc.add_paragraph('4. Write dbt tests (test-driven development)')
    doc.add_paragraph('5. Build data pipeline to pass tests')
    doc.add_paragraph('6. Create Superset dashboard and iterate')
    
    add_heading(doc, 'Feedback Instructions', 2)
    doc.add_paragraph(
        'Please edit this document directly or provide comments. Key areas for feedback:'
    )
    doc.add_paragraph('• Add/remove/modify questions each dashboard should answer', style='List Bullet')
    doc.add_paragraph('• Suggest additional charts or visualizations', style='List Bullet')
    doc.add_paragraph('• Identify missing filters or data sources', style='List Bullet')
    doc.add_paragraph('• Answer the "Questions for Nick" in each section', style='List Bullet')
    doc.add_paragraph('• Adjust priority order if needed', style='List Bullet')
    
    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = '/Users/ben/Projects/UT_Computational_NE/Neutron_OS/docs/scenarios/superset/Superset_Scenarios_For_Review.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
