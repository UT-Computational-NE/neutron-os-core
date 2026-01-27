#!/usr/bin/env python3
"""Generate Word doc questionnaires for stakeholders."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent.parent / "specs"


def create_netl_ops_questionnaire():
    """Create questionnaire for Jim and Rod at NETL Operations."""
    doc = Document()
    
    # Title
    title = doc.add_heading("Elog & Audit Requirements Questionnaire", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph("For NETL Operations Team")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Context section
    doc.add_heading("Context", level=1)
    doc.add_paragraph(
        "We are building a unified data platform for NETL reactor operations that includes "
        "an electronic logbook (Elog) system with audit capabilities. The system will support "
        "both day-to-day operations logging and regulatory compliance/audit preparation."
    )
    doc.add_paragraph(
        "Nick Luciano has reviewed our initial dashboard scenarios and identified several "
        "questions that require input from operations staff. Your answers will directly "
        "shape the Elog system design and ensure it meets your actual needs."
    )
    doc.add_paragraph(
        "The Elog will use immutable storage with cryptographic verification to ensure "
        "records cannot be tampered with after creation—a key requirement for regulatory "
        "compliance."
    )
    
    doc.add_paragraph()
    
    # Questions section
    doc.add_heading("Questions", level=1)
    
    questions = [
        {
            "num": "1",
            "question": "What categories or tags should Elog entries have?",
            "context": "We need a taxonomy for classifying log entries. Examples might include: "
                      "Startup, Shutdown, Maintenance, Safety Event, Visitor Log, Sample Insertion, "
                      "Rod Movement, etc.",
            "follow_up": "Are there required categories vs. optional tags? Should operators be able "
                        "to create new tags, or should the list be fixed?"
        },
        {
            "num": "2", 
            "question": "What constitutes a 'gap' in logging that should be flagged?",
            "context": "We want to highlight days or periods where logging may be incomplete. "
                      "This helps ensure documentation compliance.",
            "follow_up": "Is it based on time (e.g., no entries for 8+ hours during operations)? "
                        "Or based on expected events (e.g., startup logged but no shutdown)?"
        },
        {
            "num": "3",
            "question": "Should the system include export-to-PDF for audit preparation?",
            "context": "This would allow generating formatted reports of log entries for a "
                      "specified date range, filtered by category, operator, etc.",
            "follow_up": "What information must be included in an audit report? Are there "
                        "specific formats or templates currently used?"
        },
        {
            "num": "4",
            "question": "What does a typical NRC inspection request look like?",
            "context": "Understanding the format and scope of inspection requests helps us "
                      "design the right query and export capabilities.",
            "follow_up": "Can you share an example request (redacted if needed)? How much "
                        "notice do you typically receive before an inspection?"
        },
        {
            "num": "5",
            "question": "What format should evidence packages be in?",
            "context": "Evidence packages are collections of records prepared for auditors. "
                      "Options include PDF bundles, ZIP archives with source files, or "
                      "structured data exports.",
            "follow_up": "Do inspectors prefer paper printouts, electronic files, or both? "
                        "Are there specific requirements for record authenticity verification?"
        },
        {
            "num": "6",
            "question": "How far back do inspectors typically need to query?",
            "context": "This affects our data retention and indexing strategy. We can support "
                      "any retention period, but knowing typical needs helps optimize the system.",
            "follow_up": "Is it usually days, weeks, months, or years? Are there different "
                        "retention requirements for different record types?"
        },
    ]
    
    for q in questions:
        doc.add_heading(f"Question {q['num']}", level=2)
        
        # Main question in bold
        p = doc.add_paragraph()
        run = p.add_run(q['question'])
        run.bold = True
        
        # Context
        doc.add_paragraph(f"Context: {q['context']}")
        
        # Follow-up
        doc.add_paragraph(f"Follow-up: {q['follow_up']}")
        
        # Answer space
        doc.add_paragraph("Your Answer:")
        doc.add_paragraph("_" * 60)
        doc.add_paragraph("_" * 60)
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
    
    # Additional comments
    doc.add_heading("Additional Comments", level=1)
    doc.add_paragraph(
        "Please add any other requirements, concerns, or suggestions for the Elog system:"
    )
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("_" * 60)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("Thank you for your input! Please return this questionnaire to Ben.")
    doc.add_paragraph("Contact: bdb3732@utexas.edu")
    
    # Save
    output_path = OUTPUT_DIR / "Questionnaire_NETL_Operations.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_experiments_questionnaire():
    """Create questionnaire for Khiloni Shah on experiment tracking."""
    doc = Document()
    
    # Title
    title = doc.add_heading("Experiment & Sample Tracking Questionnaire", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph("For Research Operations")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Context section
    doc.add_heading("Context", level=1)
    doc.add_paragraph(
        "We are building a unified data platform for NETL that includes experiment and "
        "sample tracking capabilities. The goal is to digitize the workflow from sample "
        "preparation through irradiation, decay, counting, and analysis."
    )
    doc.add_paragraph(
        "Nick Luciano has provided an initial list of sample metadata fields (shown below). "
        "We need your input to validate this list, understand current workflows, and "
        "identify any integration needs with scheduling systems."
    )
    
    doc.add_paragraph()
    
    # Nick's proposed fields
    doc.add_heading("Proposed Sample Metadata Fields (from Nick)", level=1)
    doc.add_paragraph("Nick suggested tracking the following for each sample:")
    
    fields = [
        "Sample Name (must be unique)",
        "Sample numeric ID (assigned automatically)",
        "Chemical Composition",
        "Isotopic Composition", 
        "Density",
        "Mass",
        "Location of Irradiation (central thimble, lazy susan, etc.)",
        "Irradiation Facility (cadmium covered, etc.)",
        "Datetime of insertion",
        "Datetime of removal",
        "Decay time after removal",
        "Count Live time",
        "Total counts",
        "Total activity",
        "Activity by isotope",
        "Measurement raw data (spectra or other)"
    ]
    
    for field in fields:
        doc.add_paragraph(f"• {field}", style='List Bullet')
    
    doc.add_paragraph()
    
    # Questions section
    doc.add_heading("Questions", level=1)
    
    questions = [
        {
            "num": "1",
            "question": "How are experiments and samples currently tracked?",
            "context": "We want to understand the current workflow to ensure the new system "
                      "improves rather than disrupts existing processes.",
            "follow_up": "Is it spreadsheets, paper logs, a database, or ad-hoc? "
                        "Can you share an example of current tracking (redacted if needed)?"
        },
        {
            "num": "2",
            "question": "Are the metadata fields listed above complete and accurate?",
            "context": "Nick provided this list based on his understanding. We want to make "
                      "sure nothing is missing and the field names match your terminology.",
            "follow_up": "Are there fields to add? Remove? Rename? Are some fields optional "
                        "vs. required?"
        },
        {
            "num": "3",
            "question": "Should the system integrate with a scheduling/calendar system?",
            "context": "Nick noted that a calendar shows intent (what was planned) while the "
                      "tracking system should show what actually happened. However, linking "
                      "them could be useful.",
            "follow_up": "Do you currently use a shared calendar for scheduling reactor time? "
                        "Would it be helpful to see scheduled vs. actual in one view?"
        },
        {
            "num": "4",
            "question": "What irradiation locations and facilities should be pre-populated?",
            "context": "Nick suggested pre-populating dropdown menus for common options to "
                      "speed data entry and ensure consistency.",
            "follow_up": "Can you provide a complete list of: (a) irradiation locations "
                        "(central thimble, lazy susan positions, etc.), and (b) facility "
                        "configurations (cadmium covered, bare, etc.)?"
        },
        {
            "num": "5",
            "question": "What is the typical sample workflow from start to finish?",
            "context": "Understanding the full lifecycle helps us design a system that "
                      "supports each step.",
            "follow_up": "Walk us through a typical sample: preparation → approval → "
                        "insertion → irradiation → removal → decay → counting → analysis → "
                        "disposal. What happens at each step?"
        },
    ]
    
    for q in questions:
        doc.add_heading(f"Question {q['num']}", level=2)
        
        p = doc.add_paragraph()
        run = p.add_run(q['question'])
        run.bold = True
        
        doc.add_paragraph(f"Context: {q['context']}")
        doc.add_paragraph(f"Follow-up: {q['follow_up']}")
        
        doc.add_paragraph("Your Answer:")
        doc.add_paragraph("_" * 60)
        doc.add_paragraph("_" * 60)
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
    
    # Additional comments
    doc.add_heading("Additional Comments", level=1)
    doc.add_paragraph(
        "Please add any other requirements, concerns, or suggestions for experiment tracking:"
    )
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("_" * 60)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("Thank you for your input! Please return this questionnaire to Ben.")
    doc.add_paragraph("Contact: bdb3732@utexas.edu")
    
    output_path = OUTPUT_DIR / "Questionnaire_Experiments_Khiloni.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_cinr_concept():
    """Create CINR pre-application concept document."""
    doc = Document()
    
    # Title
    title = doc.add_heading("NEUP CINR R&D Pre-Application Concept", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Header info
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    header_data = [
        ("Submission Deadline", "January 28, 2026 (Pre-Application)"),
        ("Full Proposal (if invited)", "April 2026"),
        ("Lead Institution", "The University of Texas at Austin"),
        ("PI", "[TBD]"),
        ("Estimated Budget", "$600,000 - $800,000 over 3 years"),
    ]
    
    for i, (label, value) in enumerate(header_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Title
    doc.add_heading("Proposed Title", level=1)
    doc.add_paragraph(
        "Digital Twin Data Infrastructure for University Research Reactor Operations "
        "and Regulatory Compliance"
    )
    
    # Problem Statement
    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph(
        "University research reactors generate valuable operational, experimental, and "
        "simulation data that is currently siloed across disparate systems—spreadsheets, "
        "paper logs, simulation outputs, and instrument databases. This fragmentation:"
    )
    
    problems = [
        "Limits the ability to correlate experimental results with reactor conditions",
        "Creates compliance risk when audit trails span multiple systems",
        "Prevents effective use of digital twin simulations for operational insight",
        "Makes cross-facility data sharing and benchmarking impractical",
    ]
    for p in problems:
        doc.add_paragraph(f"• {p}", style='List Bullet')
    
    # Proposed Solution
    doc.add_heading("Proposed Solution", level=1)
    doc.add_paragraph(
        "We propose to develop an open-source data infrastructure platform that unifies "
        "research reactor data streams into a modern lakehouse architecture with:"
    )
    
    solutions = [
        "Immutable audit trails with cryptographic verification for regulatory compliance",
        "Standardized interfaces for digital twin integration (measured vs. modeled data)",
        "Real-time dashboards for operations monitoring and performance analytics",
        "Interoperability protocols enabling cross-facility data sharing",
    ]
    for s in solutions:
        doc.add_paragraph(f"• {s}", style='List Bullet')
    
    doc.add_paragraph(
        "The platform will be demonstrated at the UT Austin TRIGA reactor (NETL) with "
        "architecture designed for adoption by other university reactor facilities."
    )
    
    # Technical Approach
    doc.add_heading("Technical Approach", level=1)
    
    doc.add_heading("Data Architecture", level=2)
    doc.add_paragraph(
        "Modern lakehouse design using Apache Iceberg for time-travel queries and audit, "
        "dbt for transformation pipelines, and Dagster for orchestration. All data flows "
        "through bronze (raw) → silver (validated) → gold (analytics-ready) layers."
    )
    
    doc.add_heading("Digital Twin Integration", level=2)
    doc.add_paragraph(
        "Define a Digital Twin Provider Protocol that allows any simulation code (MPACT, "
        "OpenMC, point kinetics models) to register predictions alongside measured data. "
        "Clear labeling distinguishes measured from modeled values while enabling "
        "side-by-side visualization and validation."
    )
    
    doc.add_heading("Regulatory Compliance", level=2)
    doc.add_paragraph(
        "Implement ledger tables (append-only with bi-temporal queries) for all records "
        "requiring audit trails. Support evidence package generation for NRC inspections "
        "with cryptographic proof of record integrity."
    )
    
    doc.add_heading("Interoperability", level=2)
    doc.add_paragraph(
        "Design schema standards and APIs that enable data sharing across facilities. "
        "Align with existing nuclear ontologies (e.g., INL DIAMOND) where applicable."
    )
    
    # INL Collaboration
    doc.add_heading("Potential National Laboratory Collaboration", level=1)
    doc.add_paragraph(
        "We have identified synergies with Idaho National Laboratory's DeepLynx project, "
        "an open-source digital engineering data warehouse. Potential collaboration areas:"
    )
    
    collab = [
        "DIAMOND ontology alignment for nuclear domain vocabulary",
        "Ledger table patterns for audit compliance (DeepLynx ADR-001)",
        "MCP (Model Context Protocol) server patterns for AI-assisted data access",
        "Cross-validation of interoperability protocols",
    ]
    for c in collab:
        doc.add_paragraph(f"• {c}", style='List Bullet')
    
    doc.add_paragraph(
        "Note: This pre-application establishes foundation work. A future IRP proposal "
        "(June 2026 deadline) could formalize multi-institution collaboration with INL "
        "and other university reactor facilities."
    )
    
    # Deliverables
    doc.add_heading("Deliverables", level=1)
    
    deliverables = [
        ("Year 1", [
            "Core data platform deployed at UT Austin NETL",
            "Electronic logbook with immutable audit trail",
            "Real-time operations dashboard",
            "Digital Twin Provider Protocol specification (open source)",
        ]),
        ("Year 2", [
            "Sample/experiment tracking system",
            "MPACT shadow prediction integration",
            "Performance analytics dashboards",
            "Evidence package generator for audits",
        ]),
        ("Year 3", [
            "Cross-facility data sharing demonstration",
            "AI-assisted query tools (MCP server)",
            "Documentation and adoption guide for other facilities",
            "Final report with lessons learned",
        ]),
    ]
    
    for year, items in deliverables:
        doc.add_heading(year, level=2)
        for item in items:
            doc.add_paragraph(f"• {item}", style='List Bullet')
    
    # Relevance to NE Mission
    doc.add_heading("Relevance to DOE-NE Mission", level=1)
    doc.add_paragraph(
        "This project supports NE's mission by:"
    )
    
    relevance = [
        "Enhancing safety through better operational visibility and digital twin validation",
        "Improving regulatory compliance with modern audit infrastructure",
        "Enabling research reactor facilities to leverage advanced simulation tools",
        "Creating reusable infrastructure that benefits the broader university reactor community",
        "Developing workforce skills in modern data engineering for nuclear applications",
    ]
    for r in relevance:
        doc.add_paragraph(f"• {r}", style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("DRAFT - For internal review before submission")
    doc.add_paragraph("Contact: bdb3732@utexas.edu")
    
    output_path = OUTPUT_DIR / "CINR_PreApp_Concept_Draft.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_inl_email():
    """Create INL collaboration outreach email as Word doc."""
    doc = Document()
    
    title = doc.add_heading("INL DeepLynx Collaboration Outreach", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Email metadata
    doc.add_heading("Email Draft", level=1)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "To"
    table.rows[0].cells[1].text = "GRP-deeplynx-team@inl.gov"
    table.rows[1].cells[0].text = "Subject"
    table.rows[1].cells[1].text = "Digital Twin Data Infrastructure - Collaboration Inquiry from UT Austin"
    table.rows[2].cells[0].text = "CC"
    table.rows[2].cells[1].text = "[Your advisor if appropriate]"
    
    doc.add_paragraph()
    
    # Email body
    doc.add_heading("Message", level=1)
    
    doc.add_paragraph("Dear DeepLynx Team,")
    doc.add_paragraph()
    
    doc.add_paragraph(
        "I'm a researcher at UT Austin working on data infrastructure for nuclear systems, "
        "including our TRIGA reactor and several digital twin projects (molten salt, "
        "irradiation loops, off-gas systems). We've been studying DeepLynx Nexus and are "
        "impressed by several architectural patterns—particularly your ledger table approach "
        "(ADR-001), the MCP server implementation, and the DIAMOND ontology work."
    )
    
    doc.add_paragraph(
        "We're building an open-source data platform that unifies reactor operations data, "
        "simulation outputs, and experimental results into a modern lakehouse architecture. "
        "A key design goal is enabling digital twin predictions to be stored alongside "
        "measured data with clear provenance—so researchers can validate models against "
        "actual reactor behavior."
    )
    
    doc.add_paragraph(
        "We're also drafting an open specification for Digital Twin Interoperability that "
        "would allow different DT implementations (physics codes, ML models, hybrid approaches) "
        "to plug into shared data platforms. The goal is a lightweight protocol any research "
        "group could implement."
    )
    
    doc.add_paragraph(
        "I see potential for collaboration in several areas:"
    )
    
    areas = [
        ("DIAMOND Ontology Alignment", 
         "We'd like to use DIAMOND's domain concepts for our schema vocabulary. "
         "Could you point us to documentation on extension mechanisms?"),
        ("Ledger Table Patterns",
         "Your ADR-001 approach aligns with our regulatory compliance needs. "
         "We're implementing a similar pattern in dbt/Iceberg and would value any lessons learned."),
        ("MCP Server Patterns",
         "We're building a Python MCP server for AI-assisted data queries. "
         "Comparing notes on tool design (like your RecordTools, ProjectTools) could be valuable."),
        ("Future Joint Proposals",
         "The NEUP IRP deadline (June 2026) could be an opportunity for a multi-institution "
         "proposal on digital twin interoperability for research reactors."),
    ]
    
    for title_text, desc in areas:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        "We're not proposing to replace DeepLynx—rather, we see this as defining thin "
        "interoperability layers that let different platforms (graph-based like yours, "
        "lakehouse-based like ours, or others) exchange DT predictions and validation results."
    )
    
    doc.add_paragraph(
        "Would you be available for a 30-minute call to explore whether this aligns with "
        "INL's interests? We're also happy to share our current architecture docs for feedback."
    )
    
    doc.add_paragraph()
    doc.add_paragraph("Best regards,")
    doc.add_paragraph()
    doc.add_paragraph("[Your Name]")
    doc.add_paragraph("[Title], UT Austin")
    doc.add_paragraph("Nuclear Engineering Computational Lab")
    doc.add_paragraph("[email] | [phone]")
    
    # Attachment note
    doc.add_paragraph()
    doc.add_heading("Suggested Attachments", level=1)
    doc.add_paragraph("• DeepLynx Assessment summary (one-pager)")
    doc.add_paragraph("• Digital Twin Protocol draft (when ready)")
    doc.add_paragraph("• Architecture overview diagram")
    
    output_path = OUTPUT_DIR / "INL_Collaboration_Email_Draft.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Generating questionnaires and proposal documents...")
    create_netl_ops_questionnaire()
    create_experiments_questionnaire()
    create_cinr_concept()
    create_inl_email()
    print("\nAll documents generated successfully!")
