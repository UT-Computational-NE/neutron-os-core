#!/usr/bin/env python3
"""
Simplified Markdown to Word converter - debugging version without XML manipulations.
"""

import os
import re
import sys
import base64
import argparse
import json
import subprocess
from pathlib import Path
from io import BytesIO
from datetime import datetime
from typing import Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OxmlEl
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# =============================================================================
# DOCUMENT STYLING
# =============================================================================

def set_document_defaults(doc):
    """Set professional default fonts and spacing."""
    # Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # Headings
    for i in range(1, 7):
        heading_key = f'Heading {i}'
        if heading_key in doc.styles:
            style = doc.styles[heading_key]
            style.font.name = 'Aptos Display'
            if i == 1:
                style.font.size = Pt(20)
            elif i == 2:
                style.font.size = Pt(16)
            elif i == 3:
                style.font.size = Pt(14)
            else:
                style.font.size = Pt(12)


def set_paragraph_spacing(paragraph, after=6, before=0):
    """Set paragraph spacing - use before/after flexibility."""
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)


def add_bookmark(paragraph, bookmark_name, bookmark_id=0):
    """Add a bookmark to a paragraph for internal linking."""
    # bookmark_name should already be sanitized when passed in
    
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._r
    
    # Create bookmark start and end elements with unique ID
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    
    # Insert bookmark elements properly
    parent = tag.getparent()
    parent.insert(parent.index(tag), start)
    parent.insert(parent.index(tag) + 1, end)
    
    return bookmark_name


def add_word_toc_field(doc, toc_depth=3):
    r"""
    Add a proper Word field TOC that Office 365 recognizes and allows refreshing.
    
    This creates the standard Word TOC field code \o "1-3" \h \z \u
    which generates from actual heading styles in the document.
    
    Args:
        doc: The Word document
        toc_depth: Maximum heading level to include (default 3)
    """
    # Add TOC heading
    toc_heading = doc.add_heading('Table of Contents', level=1)
    set_paragraph_spacing(toc_heading, before=0, after=12)
    
    # Add TOC paragraph with field code
    toc_paragraph = doc.add_paragraph()
    
    # Create the field code XML
    # \o "1-3" = outline levels 1-3, \h = hyperlinks, \z = hide page numbers in web view, \u = use outline level
    field_xml = (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:instrText xml:space="preserve"> TOC \\o "1-{toc_depth}" \\h \\z \\u </w:instrText>'
        f'</w:r>'
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )
    
    # Insert field XML directly into paragraph
    for field_element_xml in field_xml.split('</w:r>')[:-1]:
        if field_element_xml.strip():
            field_element_xml += '</w:r>'
            try:
                field_element = parse_xml(field_element_xml)
                toc_paragraph._p.append(field_element)
            except Exception as e:
                print(f"Warning: Could not insert TOC field: {e}")
    
    set_paragraph_spacing(toc_paragraph, after=12)
    
    # Add a page break after TOC
    doc.add_page_break()


def add_table_of_contents(doc, headings, bookmarks, toc_depth=3):
    """
    Add a static Table of Contents with hyperlinks to the document.
    
    This is the legacy implementation that creates a manually-formatted TOC.
    For Office 365 refreshable TOC, use add_word_toc_field() instead.
    
    Args:
        doc: The Word document
        headings: List of dicts with 'level', 'text', and 'bookmark' keys
        bookmarks: Dict mapping heading text to bookmark names
        toc_depth: Maximum heading level to include (default 3)
    """
    # Add TOC heading
    toc_heading = doc.add_heading('Table of Contents', level=1)
    set_paragraph_spacing(toc_heading, before=0, after=12)
    
    # Add each heading as a TOC entry with hyperlink
    for heading in headings:
        level = heading['level']
        text = heading['text']
        bookmark = heading.get('bookmark')
        
        # Skip headings deeper than toc_depth
        if level > toc_depth:
            continue
        
        # Skip the "Table of Contents" heading itself
        if text == 'Table of Contents':
            continue
        
        # Create paragraph with indentation based on level
        p = doc.add_paragraph()
        indent = (level - 1) * 0.25  # 0.25 inch per level
        p.paragraph_format.left_indent = Inches(indent)
        set_paragraph_spacing(p, before=2, after=2)
        
        # Add hyperlink to bookmark if available
        if bookmark:
            add_hyperlink_to_paragraph(p, bookmark, text, is_internal=True)
        else:
            p.add_run(text)
    
    # Add a page break after TOC
    doc.add_page_break()


def add_hyperlink_to_paragraph(paragraph, url, text, is_internal=False):
    """Add a hyperlink to a paragraph."""
    import docx.opc.constants
    from docx.opc.part import Part
    
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    
    if is_internal:
        # For internal links, use bookmark anchor
        r_id = None
        # Create hyperlink element with anchor
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), url)
    else:
        # For external links, add relationship
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
    
    # Create new run element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    
    # Add underline  
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Add text element (must use w:t, not .text on run)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    
    # Add hyperlink to paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink


# =============================================================================
# MERMAID RENDERING
# =============================================================================

MAX_DIAGRAM_HEIGHT = 6.5  # inches
DIAGRAM_WIDTH = 6.5  # inches
DIAGRAM_CACHE = {}
CACHE_FILE = Path(__file__).parent / '.diagram_cache.json'

def load_diagram_cache():
    """Load diagram cache from disk."""
    global DIAGRAM_CACHE
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, 'r') as f:
                cache_data = json.load(f)
                DIAGRAM_CACHE = {
                    code: base64.b64decode(img_b64)
                    for code, img_b64 in cache_data.items()
                }
            print(f"Loaded {len(DIAGRAM_CACHE)} cached diagrams", flush=True)
        except Exception as e:
            print(f"Warning: Could not load diagram cache: {e}", flush=True)
            DIAGRAM_CACHE = {}

def save_diagram_cache():
    """Save diagram cache to disk."""
    try:
        cache_data = {
            code: base64.b64encode(img_bytes).decode()
            for code, img_bytes in DIAGRAM_CACHE.items()
        }
        with open(CACHE_FILE, 'w') as f:
            json.dump(cache_data, f)
        print(f"Saved {len(DIAGRAM_CACHE)} diagrams to cache", flush=True)
    except Exception as e:
        print(f"Warning: Could not save diagram cache: {e}", flush=True)

def get_diagram_dimensions(image_data: bytes) -> Optional[tuple[float, float]]:
    """Get width and height of image in inches (assuming 96 DPI)."""
    from PIL import Image
    try:
        img = Image.open(BytesIO(image_data))
        width_inches = img.width / 96
        height_inches = img.height / 96
        return (width_inches, height_inches)
    except Exception:
        return None


def render_mermaid_to_image(mermaid_code: str, timeout: int = 30, use_cache: bool = True) -> bytes | None:
    """Render Mermaid diagram to image bytes using mermaid.ink with optional caching."""
    import time
    
    # Fix common mermaid syntax issues
    sanitized_code = mermaid_code
    
    # Replace <br/> and <br> tags with spaces
    def replace_br_in_quotes(match):
        content = match.group(0)
        content = content.replace('<br/>', ' ')
        content = content.replace('<br>', ' ')
        return content
    
    sanitized_code = re.sub(r'"[^"]*"', replace_br_in_quotes, sanitized_code)
    sanitized_code = re.sub(r"'[^']*'", replace_br_in_quotes, sanitized_code)
    sanitized_code = sanitized_code.replace('<br/>', ' ')
    sanitized_code = sanitized_code.replace('<br>', ' ')
    
    # Replace emojis with text equivalents
    emoji_replacements = {
        '✅': 'ON',
        '⬜': 'OFF',
        '✔️': 'OK',
        '❌': 'X',
        '🥉': 'BRONZE',
        '🥈': 'SILVER', 
        '🥇': 'GOLD',
        '🔄': 'SYNC',
        '⚡': 'STREAM',
        '📊': 'CHART',
        '🟠': '(S)',  # Orange circle for Stale
        '🚀': '',
        '📝': '',
        '🎯': '',
        '🔥': '',
        '📈': '',
        '💎': '',
        '🎨': '',
        '🔧': '',
        '⚙️': '',
        '🧬': '',
        '🎭': '',
        '🔮': '',
        '🏆': '',
        '🌟': '',
        '💫': '',
        '⭐': '',
    }
    
    for emoji, replacement in emoji_replacements.items():
        sanitized_code = sanitized_code.replace(emoji, replacement)
    
    # Handle xychart-beta (not supported by mermaid.ink)
    if 'xychart-beta' in sanitized_code:
        lines = sanitized_code.split('\n')
        title = ''
        for line in lines:
            if 'title' in line:
                title = line.split('"')[1] if '"' in line else 'Chart'
                break
        sanitized_code = f'graph LR\n    A[{title}]\n    B[Chart visualization not supported]\n    A --> B'
    
    # Remove any remaining Unicode emojis and special characters
    emoji_pattern = r'[\U0001F300-\U0001F9FF\U00002600-\U000027BF\U0001F900-\U0001FA9F\u2000-\u206F\u2190-\u21FF]'
    sanitized_code = re.sub(emoji_pattern, '', sanitized_code)
    
    # Clean up any problematic characters but keep necessary ASCII
    safe_chars = set(range(32, 127)) | {ord('\n'), ord('\r'), ord('\t')}
    sanitized_code = ''.join(char for char in sanitized_code if ord(char) in safe_chars)
    
    # Fix labels with square brackets
    def fix_bracket_labels(text):
        pattern = r'(\w+)\["([^"]*)"\]'
        def replacer(match):
            node_id = match.group(1)
            label = match.group(2)
            return f'{node_id}["{label}"]'
        return re.sub(pattern, replacer, text)
    
    sanitized_code = fix_bracket_labels(sanitized_code)
    
    # Also handle single-quoted labels
    pattern2 = r"(\w+)\['([^']*)'\]"
    def single_quote_replacer(match):
        node_id = match.group(1)
        label = match.group(2)
        return f'{node_id}["{label}"]'
    sanitized_code = re.sub(pattern2, single_quote_replacer, sanitized_code)
    
    sanitized_code = re.sub(r'\[\s*\]', '[ ]', sanitized_code)
    
    # Check cache first
    if use_cache and mermaid_code in DIAGRAM_CACHE:
        return DIAGRAM_CACHE[mermaid_code]
    
    try:
        # Encode sanitized diagram (preserve the diagram's own theme/color settings)
        encoded = base64.urlsafe_b64encode(sanitized_code.encode()).decode()
        
        # Use white canvas background, but let diagram's own colors apply
        url = f"https://mermaid.ink/img/{encoded}?bgColor=white&width=2400"
        start = time.time()
        
        # Make request with timeout
        response = requests.get(url, timeout=timeout)
        elapsed = time.time() - start
        
        if response.status_code == 200 and response.content:
            # Cache the result
            DIAGRAM_CACHE[mermaid_code] = response.content
            return response.content
        else:
            print(f"Failed to render diagram (status={response.status_code})", flush=True)
            return None
    
    except requests.RequestException as e:
        print(f"Error rendering diagram: {e}", flush=True)
        return None
    except Exception as e:
        print(f"Unexpected error rendering diagram: {e}", flush=True)
        return None


# =============================================================================
# MARKDOWN PARSING
# =============================================================================

def parse_markdown_to_blocks(content: str) -> list:
    """Parse markdown content into structured blocks."""
    # Remove decorator patterns
    content = re.sub(r'@\w+\([^)]*\)', '', content)
    content = re.sub(r'# Review Note:.*\n', '', content, flags=re.MULTILINE)
    content = re.sub(r'^@.*$', '', content, flags=re.MULTILINE)
    
    lines = content.split('\n')
    blocks = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Skip horizontal rules
        if re.match(r'^\s*([-*_])\s*(\1\s*){2,}\s*$', line):
            i += 1
            continue
        
        # Heading
        if line.startswith('#'):
            match = re.match(r'^(#+)\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                content = match.group(2)
                blocks.append({'type': 'heading', 'level': level, 'content': content})
            i += 1
            continue
        
        # Mermaid diagram
        if line.strip().startswith('```mermaid'):
            diagram_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                diagram_lines.append(lines[i])
                i += 1
            if diagram_lines:
                blocks.append({'type': 'mermaid', 'content': '\n'.join(diagram_lines)})
            i += 1
            continue
        
        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                blocks.append({'type': 'code', 'content': '\n'.join(code_lines)})
            i += 1
            continue
        
        # Table
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'table', 'content': table_lines})
            continue
        
        # Blockquote
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                quote_lines.append(lines[i].lstrip('> ').strip())
                i += 1
            if quote_lines:
                blocks.append({'type': 'blockquote', 'content': '\n'.join(quote_lines)})
            continue
        
        # Unordered list (including nested items)
        if re.match(r'^[-*+]\s+', line):
            items = []
            while i < len(lines):
                current = lines[i]
                # Top-level bullet
                if re.match(r'^[-*+]\s+', current):
                    item = re.sub(r'^[-*+]\s+', '', current).strip()
                    if item:
                        items.append({'text': item, 'level': 0})
                    i += 1
                # Nested bullet (indented with spaces/tabs)
                elif re.match(r'^(\s+)[-*+]\s+', current):
                    indent_match = re.match(r'^(\s+)[-*+]\s+', current)
                    indent_len = len(indent_match.group(1))
                    level = max(1, indent_len // 2)  # 2-3 spaces = level 1, 4-5 = level 2, etc.
                    item = re.sub(r'^\s+[-*+]\s+', '', current).strip()
                    if item:
                        items.append({'text': item, 'level': level})
                    i += 1
                else:
                    break
            if items:
                blocks.append({'type': 'list', 'list_type': 'bullet', 'content': items})
            continue
        
        # Ordered list (including nested items)
        if re.match(r'^\d+\.\s+', line):
            items = []
            while i < len(lines):
                current = lines[i]
                # Top-level numbered
                if re.match(r'^\d+\.\s+', current):
                    item = re.sub(r'^\d+\.\s+', '', current).strip()
                    if item:
                        items.append({'text': item, 'level': 0})
                    i += 1
                # Nested bullet under numbered (indented with spaces/tabs)
                elif re.match(r'^(\s+)[-*+]\s+', current):
                    indent_match = re.match(r'^(\s+)[-*+]\s+', current)
                    indent_len = len(indent_match.group(1))
                    level = max(1, indent_len // 2)
                    item = re.sub(r'^\s+[-*+]\s+', '', current).strip()
                    if item:
                        items.append({'text': item, 'level': level})
                    i += 1
                # Nested numbered under numbered (indented)
                elif re.match(r'^(\s+)\d+\.\s+', current):
                    indent_match = re.match(r'^(\s+)\d+\.\s+', current)
                    indent_len = len(indent_match.group(1))
                    level = max(1, indent_len // 2)
                    item = re.sub(r'^\s+\d+\.\s+', '', current).strip()
                    if item:
                        items.append({'text': item, 'level': level})
                    i += 1
                else:
                    break
            if items:
                blocks.append({'type': 'list', 'list_type': 'number', 'content': items})
            continue
        
        # Regular paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip():
            if any(lines[i].startswith(p) for p in ['#', '```', '>', '|']):
                break
            if re.match(r'^([-*+]|\d+\.)\s+', lines[i]):
                break
            para_lines.append(lines[i])
            i += 1
        
        if para_lines:
            blocks.append({'type': 'paragraph', 'content': ' '.join(para_lines)})
        
        if not para_lines:
            i += 1
    
    return blocks


def process_paragraph_links(paragraph, bookmarks=None):
    """Process any pending hyperlinks in a paragraph after text is added."""
    if not hasattr(paragraph, '_pending_links'):
        return
    
    # Process each pending link
    for link_info in paragraph._pending_links:
        # Find and replace the placeholder with actual hyperlink
        for run in paragraph.runs:
            if link_info['placeholder'] in run.text:
                # Clear the placeholder text
                run.text = run.text.replace(link_info['placeholder'], '')
                
                # Get the position of this run
                run_elem = run._element
                p_elem = paragraph._p
                run_index = p_elem.index(run_elem)
                
                # Create the hyperlink
                if link_info['is_internal']:
                    # Internal link to bookmark
                    anchor_text = link_info['url'][1:]  # Remove #
                    # Try multiple variations to find the bookmark
                    anchor_key = anchor_text.lower()
                    
                    # Also try normalized form (replacing hyphens with spaces)
                    anchor_alt = anchor_text.replace('-', ' ').lower()
                    
                    # Check if bookmark exists
                    bookmark_name = None
                    if bookmarks:
                        if anchor_key in bookmarks:
                            bookmark_name = bookmarks[anchor_key]
                        elif anchor_alt in bookmarks:
                            bookmark_name = bookmarks[anchor_alt]
                    
                    if bookmark_name:
                        add_hyperlink_to_paragraph(paragraph, bookmark_name, link_info['text'], is_internal=True)
                    else:
                        # If bookmark doesn't exist, just add formatted text
                        new_run = paragraph.add_run(link_info['text'])
                        new_run.font.color.rgb = RGBColor(0, 0, 255)
                        new_run.underline = True
                else:
                    # External link
                    # For relative paths to other documents, keep them as-is for now
                    # They can be updated to permanent URLs once documents are published
                    add_hyperlink_to_paragraph(paragraph, link_info['url'], link_info['text'], is_internal=False)
                break
    
    # Clean up
    delattr(paragraph, '_pending_links')


def apply_inline_formatting(paragraph, text: str, bookmarks=None):
    """Apply inline markdown formatting to paragraph text."""
    if not text:
        return
    
    # Regex for markdown formatting - process code blocks FIRST to protect underscores
    # Then handle other formatting
    pattern = re.compile(
        r'(`[^`]+`)|'  # `code` - match this first to protect content
        r'(\*\*(.+?)\*\*)|'  # **bold**
        r'(__(.+?)__)|'      # __bold__
        r'(?<![a-zA-Z0-9_])\*(?![*\s])([^*]+)(?<![*\s])\*(?![a-zA-Z0-9_])|'  # *italic* - not inside words
        r'(?<![a-zA-Z0-9])_(?![_\s])([^_]+)(?<![_\s])_(?![a-zA-Z0-9])|'  # _italic_ - not inside words
        r'(\[(.+?)\]\((.+?)\))'  # [link](url)
    )
    
    last_end = 0
    
    for match in pattern.finditer(text):
        # Add plain text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        matched_text = match.group(0)
        
        # Code: `text` - handle first to protect underscores
        if matched_text.startswith('`'):
            content = matched_text[1:-1]  # Remove backticks
            run = paragraph.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        # Bold: **text** or __text__
        elif matched_text.startswith('**'):
            content = match.group(3)
            run = paragraph.add_run(content)
            run.bold = True
        elif matched_text.startswith('__'):
            content = match.group(5)
            run = paragraph.add_run(content)
            run.bold = True
        
        # Italic: *text* or _text_ (but not inside words)
        elif matched_text.startswith('*') and not matched_text.startswith('**'):
            content = match.group(6)
            if content:
                run = paragraph.add_run(content)
                run.italic = True
        elif matched_text.startswith('_') and not matched_text.startswith('__'):
            content = match.group(7)
            if content:
                run = paragraph.add_run(content)
                run.italic = True
        
        # Link: [text](url)
        elif matched_text.startswith('['):
            content = match.group(9)
            url = match.group(10)
            if content and url:
                # Store the link info for processing after we finish the plain text
                if not hasattr(paragraph, '_pending_links'):
                    paragraph._pending_links = []
                # Mark position where link should be inserted
                placeholder = f"__LINK_{len(paragraph._pending_links)}__"
                paragraph.add_run(placeholder)
                paragraph._pending_links.append({
                    'text': content,
                    'url': url,
                    'placeholder': placeholder,
                    'is_internal': url.startswith('#')
                })
        
        last_end = match.end()
    
    # Add remaining text after last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])
    
    # Process any pending links
    process_paragraph_links(paragraph, bookmarks)


def clean_inline_formatting(text: str) -> str:
    """Remove markdown formatting while preserving content."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text.strip()


# =============================================================================
# DOCUMENT GENERATION
# =============================================================================

def add_table_to_doc(doc, headers: list[str], rows: list[list[str]], bookmarks=None):
    """Add a formatted table."""
    if not headers:
        return
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Set vertical alignment for all cells to top
    def set_cell_vertical_alignment(cell, alignment='top'):
        """Set vertical alignment of a table cell."""
        tcPr = cell._element.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), alignment)
        tcPr.append(tcVAlign)
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_vertical_alignment(cell, 'top')
        # Clear default text and apply formatting
        cell.text = ''
        para = cell.paragraphs[0]
        apply_inline_formatting(para, header, bookmarks)
        # Make all runs in header bold
        for run in para.runs:
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        set_paragraph_spacing(para, after=6)
    
    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row):
                set_cell_vertical_alignment(row[i], 'top')
                # Clear default text and apply formatting
                row[i].text = ''
                para = row[i].paragraphs[0]
                apply_inline_formatting(para, cell_data, bookmarks)
                # Set font for all runs
                for run in para.runs:
                    if not run.font.name:  # Don't override code font
                        run.font.name = 'Calibri'
                set_paragraph_spacing(para, after=4)
    
    # Space after table
    doc.add_paragraph()


def add_mermaid_to_doc(doc, mermaid_code: str):
    """Add a Mermaid diagram - simplified version without XML manipulation."""
    image_data = render_mermaid_to_image(mermaid_code)
    
    if image_data:
        image_stream = BytesIO(image_data)
        try:
            # Get dimensions for validation and scaling
            dims = get_diagram_dimensions(image_data)
            
            max_width = Inches(6.5)
            max_height = Inches(MAX_DIAGRAM_HEIGHT)
            
            # Calculate actual dimensions to use
            final_width = max_width
            final_height = None
            
            if dims:
                width_inches, height_inches = dims
                aspect_ratio = width_inches / height_inches
                
                # Check if we need to scale down due to height constraint
                if height_inches > MAX_DIAGRAM_HEIGHT:
                    print(f"    [Scaling: Diagram height {height_inches:.1f}\" exceeds {MAX_DIAGRAM_HEIGHT}\" limit - scaling down]", flush=True)
                    final_height = max_height
                    final_width = Inches(MAX_DIAGRAM_HEIGHT * aspect_ratio)
                    
                    if final_width > max_width:
                        final_width = max_width
                        final_height = Inches(6.5 / aspect_ratio)
                
                else:
                    potential_height = 6.5 / aspect_ratio
                    if potential_height > MAX_DIAGRAM_HEIGHT:
                        final_height = max_height
                        final_width = Inches(MAX_DIAGRAM_HEIGHT * aspect_ratio)
            
            # Add paragraph for diagram - NO XML MANIPULATION
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add the image
            if final_height:
                run = p.add_run()
                run.add_picture(image_stream, width=final_width, height=final_height)
            else:
                run = p.add_run()
                run.add_picture(image_stream, width=final_width)
            
            # Set spacing after diagram (instead of adding blank paragraph)
            set_paragraph_spacing(p, after=12)
            
        except Exception as e:
            import traceback
            print(f"  Error adding diagram to document: {e}", flush=True)
            print(f"  Traceback: {traceback.format_exc()}", flush=True)
            p = doc.add_paragraph("⚠️ Diagram could not be embedded")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, after=12)
    else:
        p = doc.add_paragraph("⚠️ Diagram rendering failed")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=12)


def markdown_to_docx(input_file: str, output_file: str, pregenerated_images: dict = None, verbose: bool = True, generate_bookmarks: bool = False, generate_toc: bool = False):
    """Convert markdown file to Word document.
    
    Args:
        input_file: Path to markdown file
        output_file: Path to output Word document
        pregenerated_images: Dict of pregenerated mermaid images
        verbose: Print progress messages
        generate_bookmarks: Whether to create Word bookmarks for headings (default: False)
        generate_toc: Whether to generate Table of Contents (default: False)
                     When True, uses proper Word field code that Office 365 recognizes as refreshable
    """
    # Read markdown content
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse into blocks
    blocks = parse_markdown_to_blocks(content)
    
    # Create document
    doc = Document()
    set_document_defaults(doc)
    
    # First pass: pre-register all heading bookmarks and build TOC entries
    bookmarks = {}  # Track bookmarks for internal links
    toc_entries = []  # List of headings for TOC
    for block in blocks:
        if block['type'] == 'heading':
            text = block['content']
            plain_text = clean_inline_formatting(text)
            # Generate the bookmark name that will be used
            bookmark_name = re.sub(r'[^a-zA-Z0-9_-]', '_', plain_text)
            if bookmark_name and not bookmark_name[0].isalpha():
                bookmark_name = 'bm_' + bookmark_name
            bookmark_name = bookmark_name[:40]
            
            # Store with multiple key formats for flexible matching
            bookmarks[plain_text.lower()] = bookmark_name
            # Also store with normalized form (like "1-executive-summary" format)
            normalized_key = re.sub(r'[^\w\s-]', '', plain_text.lower())
            normalized_key = re.sub(r'[-\s]+', '-', normalized_key).strip('-')
            bookmarks[normalized_key] = bookmark_name
            # Bookmark pre-registered
            
            # Add to TOC entries
            if generate_toc:
                toc_entries.append({
                    'level': block['level'],
                    'text': plain_text,
                    'bookmark': bookmark_name
                })
    
    # Generate Table of Contents if requested
    if generate_toc:
        # Use proper Word field code TOC that Office 365 recognizes and can refresh
        add_word_toc_field(doc, toc_depth=3)
    
    # Second pass: process blocks with bookmarks available
    diagram_count = 0
    block_count = 0
    bookmark_id = 0
    
    for i, block in enumerate(blocks):
        block_count += 1
        block_type = block['type']
        
        if verbose:
            if block_type == 'mermaid':
                diagram_count += 1
                print(f"  [{block_count}/{len(blocks)}] Diagram {diagram_count}...", end='', flush=True)
            else:
                print(f"  [{block_count}/{len(blocks)}] {block_type.capitalize()}", flush=True)
        
        if block_type == 'heading':
            level = block['level']
            text = block['content']
            
            if level <= 6:
                # Add heading without text first
                p = doc.add_heading('', level=level)
                # Apply inline formatting to the heading
                apply_inline_formatting(p, text, bookmarks)
                
                # Check if next block is a diagram - if so, use minimal spacing
                # This keeps headings connected to their diagrams
                next_is_diagram = (i + 1 < len(blocks) and 
                                 blocks[i + 1]['type'] == 'mermaid')
                
                # Headings should have MORE space before them than after
                # This visually associates them with the content below
                if next_is_diagram:
                    # Minimal spacing when followed by diagram
                    set_paragraph_spacing(p, before=12, after=2)
                else:
                    # More space before heading, less after (to group with content below)
                    set_paragraph_spacing(p, before=12, after=4)
                
                # Add bookmark for this heading
                # Use the plain text version for the bookmark name
                plain_text = clean_inline_formatting(text)
                # Get the pre-registered bookmark name
                normalized_key = re.sub(r'[^\w\s-]', '', plain_text.lower())
                normalized_key = re.sub(r'[-\s]+', '-', normalized_key).strip('-')
                if normalized_key in bookmarks:
                    bookmark_name = bookmarks[normalized_key]
                else:
                    bookmark_name = bookmarks.get(plain_text.lower(), None)
                
                if generate_bookmarks and bookmark_name:
                    # Actually create the bookmark in the document
                    add_bookmark(p, bookmark_name, bookmark_id)
                    bookmark_id += 1
        
        elif block_type == 'paragraph':
            p = doc.add_paragraph()
            apply_inline_formatting(p, block['content'], bookmarks)
            set_paragraph_spacing(p)
        
        elif block_type == 'blockquote':
            p = doc.add_paragraph(style='Quote')
            apply_inline_formatting(p, block['content'], bookmarks)
            set_paragraph_spacing(p)
        
        elif block_type == 'code':
            p = doc.add_paragraph()
            run = p.add_run(block['content'])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            set_paragraph_spacing(p)
        
        elif block_type == 'list':
            list_type = block.get('list_type', 'bullet')
            
            for idx, item in enumerate(block['content']):
                # Handle both old format (string) and new format (dict with text/level)
                if isinstance(item, dict):
                    item_text = item['text']
                    item_level = item.get('level', 0)
                else:
                    item_text = item
                    item_level = 0
                
                # Create paragraph without Word list styles (which have unpredictable indent)
                p = doc.add_paragraph()
                
                # Consistent hanging indent for all list items
                base_indent = Inches(0.25)      # Overall list indent
                level_indent = Inches(0.25 * item_level)  # Additional indent per nesting level (reduced from 0.5)
                p.paragraph_format.left_indent = base_indent + level_indent
                p.paragraph_format.first_line_indent = Inches(-0.25)  # Negative hangindent pulls bullet back
                
                # Add bullet or number
                if list_type == 'bullet':
                    if item_level == 0:
                        bullet_char = '• '  # Standard bullet
                    else:
                        bullet_char = '◦ '  # Nested bullet (hollow circle)
                else:  # Numbered list
                    # For simplicity, use dashes for nested numbered items
                    if item_level == 0:
                        bullet_char = f'{item.get("number", "•")} '  # Would need to track number in parser
                    else:
                        bullet_char = '- '  # Nested item under numbered
                
                bullet_run = p.add_run(bullet_char)
                bullet_run.font.name = 'Calibri'
                bullet_run.font.size = Pt(11)
                
                # Add the item text
                apply_inline_formatting(p, item_text, bookmarks)
                
                # Spacing: more before first item, less between items
                if idx == 0:
                    set_paragraph_spacing(p, before=6, after=3)  # Space before list
                else:
                    set_paragraph_spacing(p, before=0, after=3)  # Space between items
        
        elif block_type == 'table':
            table_data = block['content']
            if len(table_data) >= 2:
                headers = [cell.strip() for cell in table_data[0].split('|')[1:-1]]
                rows = []
                for row in table_data[2:]:
                    cells = [cell.strip() for cell in row.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                if headers and rows:
                    add_table_to_doc(doc, headers, rows, bookmarks)
        
        elif block_type == 'mermaid':
            # Use pregenerated image if available
            if pregenerated_images and block['content'] in pregenerated_images:
                DIAGRAM_CACHE[block['content']] = pregenerated_images[block['content']]
            
            add_mermaid_to_doc(doc, block['content'])
            
            if verbose:
                print("", flush=True)  # New line after diagram status
    
    # Save document
    if verbose:
        print(f"  Saving document...", end='', flush=True)
    
    doc.save(output_file)
    
    if verbose:
        print(" Done!", flush=True)
        print(f"\nGenerated: {os.path.basename(output_file)}")
        print(f"  Processed {len(blocks)} blocks ({diagram_count} diagrams)")
    
    return doc


def main():
    parser = argparse.ArgumentParser(description='Convert markdown to Word document')
    parser.add_argument('source', help='Input markdown file')
    parser.add_argument('output', nargs='?', help='Output Word document')
    parser.add_argument('--pregenerate', action='store_true', help='Pre-generate all diagrams')
    parser.add_argument('--bookmarks', action='store_true', help='Generate Word bookmarks for headings (default: disabled)')
    parser.add_argument('--toc', action='store_true', help='Generate Table of Contents using Word field code (default: disabled)')
    
    args = parser.parse_args()
    
    # Determine output filename
    if args.output:
        output_file = args.output
    else:
        input_path = Path(args.source)
        output_file = input_path.with_suffix('.docx')
    
    # Load cache
    load_diagram_cache()
    
    # Pregenerate diagrams if requested
    pregenerated = {}
    if args.pregenerate:
        with open(args.source, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Find all mermaid blocks
        mermaid_pattern = r'```mermaid\n(.*?)\n```'
        diagrams = re.findall(mermaid_pattern, content, re.DOTALL)
        
        if diagrams:
            print(f"\nPregenerating {len(diagrams)} diagram(s) asynchronously...", flush=True)
            
            def render_single(idx_code):
                idx, code = idx_code
                try:
                    result = render_mermaid_to_image(code, use_cache=False)
                    if result:
                        print(f"  ✓ Diagram {idx}/{len(diagrams)} rendered", flush=True)
                        return (code, result)
                    else:
                        print(f"  ✗ Diagram {idx}/{len(diagrams)} failed", flush=True)
                        return None
                except Exception as e:
                    print(f"  ✗ Diagram {idx}/{len(diagrams)} error: {e}", flush=True)
                    return None
            
            with ThreadPoolExecutor(max_workers=5) as executor:
                futures = [executor.submit(render_single, (i+1, d)) for i, d in enumerate(diagrams)]
                
                for future in as_completed(futures):
                    result = future.result()
                    if result:
                        code, image_data = result
                        pregenerated[code] = image_data
            
            # Update cache with pregenerated images
            DIAGRAM_CACHE.update(pregenerated)
            save_diagram_cache()
            
            success_count = len(pregenerated)
            fail_count = len(diagrams) - success_count
            print(f"Pregeneration complete: {success_count} rendered, {fail_count} failed\n", flush=True)
    
    # Convert markdown to Word
    markdown_to_docx(args.source, str(output_file), pregenerated, generate_bookmarks=args.bookmarks, generate_toc=args.toc)
    
    # Save cache
    save_diagram_cache()


if __name__ == '__main__':
    main()