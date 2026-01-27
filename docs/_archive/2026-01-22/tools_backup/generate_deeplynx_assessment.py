#!/usr/bin/env python3
"""
Generate DeepLynx Assessment Word document for UT Nuclear Engineering.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_document():
    doc = Document()
    
    # Set narrow margins for one-page fit
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('INL DeepLynx Nexus Assessment', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('UT Nuclear Engineering Digital Twin Infrastructure')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph(f'Date: January 15, 2026 | Classification: Internal Planning')
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'Idaho National Laboratory\'s DeepLynx Nexus is an open-source digital engineering data warehouse '
        'that stores data in a graph format following user-defined ontologies. After technical review, '
        'we recommend cherry-picking specific patterns rather than adopting the platform wholesale.'
    )
    
    # Two-column comparison table
    doc.add_heading('Assessment Summary', level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '✅ Worth Adopting'
    hdr_cells[1].text = '❌ Skip (Outdated for LLM Era)'
    
    # Make headers bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    row = table.add_row().cells
    row[0].text = (
        '• Ledger Table Pattern\n'
        '  Denormalized audit snapshots\n'
        '  Self-contained history\n\n'
        '• MCP Server Pattern\n'
        '  AI tool integration for agents\n'
        '  Query/create via LLM tools\n\n'
        '• Ontology-as-Code\n'
        '  YAML domain model definitions\n'
        '  Generate validators from schema'
    )
    row[1].text = (
        '• GraphQL Interface\n'
        '  LLMs generate SQL better\n\n'
        '• Manual Type Mapping\n'
        '  LLM-assisted inference is superior\n\n'
        '• Enterprise Tool Connectors\n'
        '  Not relevant for research\n\n'
        '• Graph-First Data Model\n'
        '  SQL more ML/pandas-friendly'
    )
    
    # Recommendation
    doc.add_heading('Recommendation', level=2)
    doc.add_paragraph(
        'Cherry-pick patterns without deploying DeepLynx:'
    )
    
    rec_table = doc.add_table(rows=4, cols=2)
    rec_table.style = 'Table Grid'
    
    recommendations = [
        ('Ledger Tables', 'Implement in dbt for audit trails'),
        ('Ontology YAML', 'Create our own format for domain models'),
        ('MCP Server', 'Build Python equivalent for AI agents'),
        ('GraphQL', 'Skip — stay with SQL for analytics'),
    ]
    
    for i, (pattern, action) in enumerate(recommendations):
        rec_table.rows[i].cells[0].text = pattern
        rec_table.rows[i].cells[1].text = action
    
    # Collaboration
    doc.add_heading('Potential INL Collaboration', level=2)
    doc.add_paragraph(
        'If collaboration makes sense, explore informally:'
    )
    doc.add_paragraph('• Nuclear ontology standardization (DIAMOND compatibility)', style='List Bullet')
    doc.add_paragraph('• MCP tool interface specifications', style='List Bullet')
    doc.add_paragraph('• Regulatory compliance patterns for digital twins', style='List Bullet')
    
    doc.add_paragraph(
        'Avoid: Merging codebases, shared infrastructure, or adopting DeepLynx as primary data store.',
        style='Intense Quote'
    )
    
    # Links
    doc.add_heading('References', level=2)
    p = doc.add_paragraph()
    p.add_run('GitHub: ').bold = True
    p.add_run('github.com/idaholab/DeepLynx')
    p.add_run(' | ')
    p.add_run('Product: ').bold = True
    p.add_run('inlsoftware.inl.gov/product/deep-lynx')
    
    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = '/Users/ben/Projects/UT_Computational_NE/Neutron_OS/docs/specs/DeepLynx_Assessment_UT_NE.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
