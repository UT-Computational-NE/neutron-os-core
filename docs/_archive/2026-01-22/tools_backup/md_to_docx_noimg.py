#!/usr/bin/env python3
"""Convert markdown to Word document without images for debugging."""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_simple(md_path):
    """Simple markdown parser that skips mermaid blocks."""
    with open(md_path, 'r') as f:
        content = f.read()
    
    blocks = []
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip mermaid blocks
        if line.strip().startswith('```mermaid'):
            while i < len(lines) and not lines[i].strip().endswith('```'):
                i += 1
            i += 1
            blocks.append({'type': 'mermaid_skip', 'content': '[Diagram removed for debugging]'})
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            content = line.lstrip('#').strip()
            blocks.append({'type': 'heading', 'level': level, 'content': content})
        
        # Tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            i -= 1
            blocks.append({'type': 'table', 'content': '\n'.join(table_lines)})
        
        # Paragraphs
        elif line.strip():
            blocks.append({'type': 'paragraph', 'content': line})
        
        i += 1
    
    return blocks

def convert_simple(md_path, output_path):
    """Convert markdown to docx without images."""
    doc = Document()
    blocks = parse_markdown_simple(md_path)
    
    for block in blocks:
        if block['type'] == 'heading':
            doc.add_heading(block['content'], level=min(block['level'], 9))
        elif block['type'] == 'paragraph':
            doc.add_paragraph(block['content'])
        elif block['type'] == 'mermaid_skip':
            p = doc.add_paragraph(block['content'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        elif block['type'] == 'table':
            # Simple table parsing
            lines = block['content'].split('\n')
            if len(lines) >= 3:
                headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
                rows = []
                for line in lines[2:]:
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    rows.append(cells)
                
                if headers:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    for i, header in enumerate(headers):
                        table.rows[0].cells[i].text = header
                    for row_data in rows:
                        row = table.add_row()
                        for i, cell_data in enumerate(row_data):
                            if i < len(row.cells):
                                row.cells[i].text = cell_data
    
    doc.save(output_path)
    print(f"Created {output_path} (no images)")
    print(f"Processed {len(blocks)} blocks")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx_noimg.py input.md output.docx")
        sys.exit(1)
    
    convert_simple(sys.argv[1], sys.argv[2])